from GeneralApp.models                  import Service, Provider, Property, TRANSFER_TYPE_CHOICES, Unit, Hotel
from SalesApp                           import serializers, models, querysets, views
from GeneralApp.utils                   import CustomValidation
from rest_framework                     import viewsets, permissions, status
from django.template.loader             import get_template
from django.http                        import HttpResponse,JsonResponse
from rest_framework.response            import Response
from datetime                           import time, date, datetime, timezone, timedelta
from io                                 import BytesIO
from xhtml2pdf                          import pisa
from django.conf                        import settings
import xlsxwriter
from docx                               import Document
from docx.shared                        import Inches, Pt
from docx.enum.text                     import WD_ALIGN_PARAGRAPH
from docx.enum.section                  import WD_ORIENT
from GeneralApp.serializers             import RepresentativeSerializer, UserExtensionSerializer, SaleTypeSerializer, ProviderSerializer
from GeneralApp.models                  import Representative, ExchangeRate, SaleType, Service, Provider, Group, AvailabilityGroup, UserExtension
from operacionesVP import serverconfig
from django.db import connection



def sale_subtotal(sale, with_out_tax = False):
    import locale
    locale.setlocale( locale.LC_ALL, 'en_US.UTF-8' )
    context = {}
    if sale.status == "C":
        context['adult_fee_num'] = 0
        context['child_fee_num'] = 0
        context['subtotal'] = locale.currency(0,grouping=True)
        context['subtotal_num'] = 0
        context['discount'] = locale.currency(0,grouping=True)
        context['discount_num'] = 0
        context['total'] = locale.currency(0,grouping=True)
        context['total_num'] = 0
        context['total_mn_num'] = 0
        context['total_mn'] = locale.currency(0,grouping=True)
        context['total_no_tax'] = 0
        context['total_only_tax'] = 0
        context['total_cost_num'] = 0
        context['direct_sale'] = 0
        context['direct_sale_cost'] = 0
        context['cr_crc'] = 0
        context['cr_crc_cost'] = 0
        context['rep_comission'] = 0
        context['no_comis'] = 0
        context['no_comis_cost'] = 0
        context['sup_comis'] = 0
        context['coords_comis'] = 0
    else:
        #exchange_rate = ExchangeRate.objects.filter(start_date__lte=sale.sale_date.date(),type='SALE',property=sale.property).order_by("-start_date").first()
        #coordinators_comission = models.CoordinatorsComission.objects.filter(date__lte=sale.sale_date.date(),property=sale.property).order_by("-date").first()
        #print("Comparacion {}:{}".format(exchange_rate.usd_currency,sale.usd_currency))
        service_rate_id = sale.service_rate.id
        service_rate_currency = sale.currency if sale.currency is not None else sale.service_rate.currency
        service_rate_adult_price = sale.adult_price if sale.adult_price is not None else sale.service_rate.adult_price
        service_rate_child_price = sale.child_price if sale.child_price is not None else sale.service_rate.child_price
        service_rate_exent_import_adult = sale.exent_import_adult if sale.exent_import_adult is not None else sale.service_rate.exent_import_adult
        service_rate_exent_import_child = sale.exent_import_child if sale.exent_import_child is not None else sale.service_rate.exent_import_child
        service_rate_tax = sale.tax if sale.tax is not None else sale.service_rate.tax
        service_rate_hard_rock_comission_adult = sale.hard_rock_comission_adult if sale.hard_rock_comission_adult is not None else sale.service_rate.hard_rock_comission_adult
        service_rate_hard_rock_comission_child = sale.hard_rock_comission_child if sale.hard_rock_comission_child is not None else sale.service_rate.hard_rock_comission_child
        adult_price = sale.service_fee/100
        child_price = 0
        exent_import_adult = 0
        exent_import_child = 0
        if sale.sale_service_fee is True:
            if service_rate_currency == "MN":
                adult_price = adult_price/sale.usd_currency
            elif service_rate_currency == "EURO":
                adult_price = (adult_price*sale.euro_currency)/sale.usd_currency
        else:
            adult_price = service_rate_adult_price
            child_price = service_rate_child_price
            exent_import_adult = service_rate_exent_import_adult
            exent_import_child = service_rate_exent_import_child
            if service_rate_currency == "MN":
                adult_price = adult_price/sale.usd_currency
                child_price = child_price/sale.usd_currency
                exent_import_adult = exent_import_adult/sale.usd_currency
                exent_import_child = exent_import_child/sale.usd_currency
            elif service_rate_currency == "EURO":
                adult_price = (adult_price*sale.euro_currency)/sale.usd_currency
                child_price = (child_price*sale.euro_currency)/sale.usd_currency
                exent_import_adult = (exent_import_adult*sale.euro_currency)/sale.usd_currency
                exent_import_child = (exent_import_child*sale.euro_currency)/sale.usd_currency
        discount = sale.discount
        if sale.discount_type == "percent":
            discount = (discount/ 100) * ((adult_price*sale.adults) + (child_price*sale.childs))
        discount = sale.overcharged - discount
        #encaso de pedir todo sin iva
        if with_out_tax is True and service_rate_tax > 0:
            discount = discount/((service_rate_tax/100)+1)
            adult_price = ((adult_price-exent_import_adult)/((service_rate_tax/100)+1)) + exent_import_adult
            child_price = ((child_price-exent_import_child)/((service_rate_tax/100)+1)) + exent_import_child
        sale_total = (adult_price*sale.adults) + (child_price*sale.childs)
        adult_fee = adult_price-(adult_price*(service_rate_hard_rock_comission_adult/100))
        child_fee = child_price-(child_price*(service_rate_hard_rock_comission_child/100))
        #costos sin iva
        sale_cost = (adult_fee*sale.adults) + (child_fee*sale.childs)
        sale_cost_no_tax = sale_cost
        if with_out_tax is False and service_rate_tax > 0:
            adult_price_no_tax = ((adult_price-exent_import_adult)/((service_rate_tax/100)+1)) + exent_import_adult
            child_price_no_tax = ((child_price-exent_import_child)/((service_rate_tax/100)+1)) + exent_import_child
            adult_fee_no_tax = adult_price_no_tax-(adult_price_no_tax*(service_rate_hard_rock_comission_adult/100))
            child_fee_no_tax = child_price_no_tax-(child_price_no_tax*(service_rate_hard_rock_comission_child/100))
            sale_cost_no_tax = (adult_fee_no_tax*sale.adults) + (child_fee_no_tax*sale.childs)
        total = sale_total + discount
        context['adult_fee_num'] = adult_fee
        context['child_fee_num'] = child_fee
        context['subtotal_num'] = sale_total
        context['total_cost_num'] = sale_cost
        context['total_cost_num_no_tax'] = sale_cost_no_tax
        context['discount_num'] = discount
        context['total_num'] = total
        context['direct_sale'] = 0
        context['direct_sale_cost'] = 0
        context['cr_crc'] = 0
        context['cr_crc_cost'] = 0
        context['rep_comission'] = 0
        context['no_comis'] = 0
        context['no_comis_cost'] = 0
        context['sup_comis'] = 0
        context['coords_comis'] = 0
        for sale_payment in sale.sale_payments.all():
            amount = sale_payment.amount
            if sale_payment.payment_method.currency == "MN":
                amount = amount/sale.usd_currency
            elif sale_payment.payment_method.currency == "EURO":
                amount = (amount*sale.euro_currency)/sale.usd_currency
            if "Directa" in sale_payment.payment_method.payment_type.name:
                context['direct_sale'] += amount
                if total != 0:
                    amount_no_tax = amount/((service_rate_tax/100)+1)
                    percent_payment_cost = amount_no_tax/total
                    payment_cost = percent_payment_cost * sale_cost
                    context['direct_sale_cost'] += payment_cost
                if sale.coordinators_comission is not None:
                    coords_comis = (amount/((service_rate_tax/100)+1))*(sale.coordinators_comission/100)
                    context['coords_comis'] += coords_comis
                    context['sup_comis'] += coords_comis
            if "RC / RCC" in sale_payment.payment_method.payment_type.name:
                context['cr_crc'] += amount
                if total != 0:
                    amount_no_tax = amount/((service_rate_tax/100)+1)
                    percent_payment_cost = amount_no_tax/total
                    payment_cost = percent_payment_cost * sale_cost
                    context['cr_crc_cost'] += payment_cost
            if sale_payment.payment_method.payment_type.is_commissionable is True:
                service_rate_comissions = models.ServiceRateComission.objects.filter(service_rate__id=service_rate_id,payment_type=sale_payment.payment_method.payment_type)
                for service_rate_comission in service_rate_comissions:
                    if service_rate_comission.comission > 0:
                        comission = (amount/((service_rate_tax/100)+1))*(service_rate_comission.comission/100)
                        context['rep_comission'] += comission
                        if "RC / RCC" in sale_payment.payment_method.payment_type.name:
                            context['coords_comis'] += (3 / 100) * comission
                            context['sup_comis'] += (1 / 100) * comission
            else:
                context['no_comis'] += amount
                if total != 0:
                    amount_no_tax = amount/((service_rate_tax/100)+1)
                    percent_payment_cost = amount_no_tax/total
                    payment_cost = percent_payment_cost * sale_cost
                    context['no_comis_cost'] += payment_cost
        context['subtotal_num'] = round(context['subtotal_num'] * -1  if sale.status == "R" else context['subtotal_num'], 2)
        context['subtotal'] = locale.currency(context['subtotal_num'],grouping=True)
        context['total_cost_num'] = round(context['total_cost_num'] * -1  if sale.status == "R" else context['total_cost_num'], 2)
        context['total_cost_num_no_tax'] = round(context['total_cost_num_no_tax'] * -1  if sale.status == "R" else context['total_cost_num_no_tax'], 2)
        context['discount_num'] = round(context['discount_num'] * -1  if sale.status == "R" else context['discount_num'], 2)
        context['discount'] = locale.currency(abs(context['discount_num']),grouping=True)
        context['total_mn_num'] = round((context['total_num'] * -1)*sale.usd_currency  if sale.status == "R" else context['total_num']*sale.usd_currency, 2)
        context['total_mn'] = locale.currency(context['total_mn_num'],grouping=True)
        context['total_num'] = round(context['total_num'] * -1  if sale.status == "R" else context['total_num'], 2)
        context['total'] = locale.currency(context['total_num'],grouping=True)
        if service_rate_tax > 0:
            total_no_tax = context['total_num']/((service_rate_tax/100)+1)
            context['total_no_tax'] = round(total_no_tax, 2)
            total_only_tax = total_no_tax*(service_rate_tax/100)
            context['total_only_tax'] = round(total_only_tax, 2)
        else:
            context['total_no_tax'] = 0
            context['total_only_tax'] = 0
        context['direct_sale'] = round(context['direct_sale'] * -1  if sale.status == "R" else context['direct_sale'], 2)
        context['direct_sale_cost'] = round(context['direct_sale_cost'] * -1  if sale.status == "R" else context['direct_sale_cost'], 2)
        context['cr_crc'] = round(context['cr_crc'] * -1  if sale.status == "R" else context['cr_crc'], 2)
        context['cr_crc_cost'] = round(context['cr_crc_cost'] * -1  if sale.status == "R" else context['cr_crc_cost'], 2)
        context['rep_comission'] = round(context['rep_comission'] * -1  if sale.status == "R" else context['rep_comission'], 2)
        context['no_comis'] = round(context['no_comis'] * -1  if sale.status == "R" else context['no_comis'], 2)
        context['no_comis_cost'] = round(context['no_comis_cost'] * -1  if sale.status == "R" else context['no_comis_cost'], 2)
        context['sup_comis'] = round(context['sup_comis'] * -1  if sale.status == "R" else context['sup_comis'], 2)
        context['coords_comis'] = round(context['coords_comis'] * -1  if sale.status == "R" else context['coords_comis'], 2)
    return context

def daterange(start_date, due_date):
    due_date = due_date + timedelta(1)
    for n in range(int((due_date - start_date).days)):
        yield start_date + timedelta(n)

class SaleListController():
    def __init__(self, request, start_date, due_date, property, type, format, sale_types, hotels, services, providers, representatives, with_out_tax, date_filter, just_import, summary):
        from GeneralApp.models import ExchangeRate
        self.request = request
        self.start_date = start_date
        self.due_date = due_date
        self.property = property
        self.type = type
        self.format = format
        self.sale_types = sale_types
        self.representatives = representatives
        self.hotels = hotels
        self.services = services
        self.providers = providers
        self.with_out_tax = with_out_tax
        self.date_filter = date_filter
        self.just_import = just_import
        self.summary = summary

    def filters(self):
        sales = models.Sale.objects.tableOptimization()
        sales = sales.exclude(status__in=["P","B"])
        if self.date_filter == "xFechaVenta":
            sales = sales.saleFilterSaleDateRange(
                self.start_date,self.due_date,self.property
            )
        elif self.date_filter == "xFechaServicio":
            sales = sales.saleFilterServiceDateRange(
                self.start_date,self.due_date,self.property
            )
        else:
            raise CustomValidation("No hay filtro para busqueda de fechas", 'date_filter', status.HTTP_400_BAD_REQUEST)
        if len(self.sale_types) > 0:
            sales = sales.saleFilterBySaleTypes(self.sale_types)
        if len(self.providers) > 0:
            sales = sales.saleFilterByProviders(self.providers)
        if self.type == "report_by_payment_method" and self.request.user.has_perm('SalesApp.access_report_by_payment_method_sales') and self.request.user.is_superuser is False:
            user_extension = UserExtension.objects.get(user=self.request.user)
            if user_extension.representative is not None:
                sales = sales.saleFilterByRepresentatives([user_extension.representative.id])
            else:
                raise CustomValidation("No tiene representante para ver el reporte", 'representative', status.HTTP_400_BAD_REQUEST)
        elif len(self.representatives) > 0:
            sales = sales.saleFilterByRepresentatives(self.representatives)
        if len(self.hotels) > 0:
            sales = sales.saleFilterByHotels(self.hotels)
        if len(self.services) > 0:
            sales = sales.saleFilterByServices(self.services)
        return sales
    
    def get_context(self):
        sales = self.filters()
        if self.type == "by_representative":
            return self.report_detail_by_representative(sales)
        if self.type == "summary_by_representatives":
            return self.report_summary_by_representative(sales)
        if self.type == "summary_by_sales_types_services":
            return self.report_summary_by_sales_types_services(sales)
        if self.type == "summary_by_services":
            return self.report_summary_by_services(sales)
        if self.type == "summary_by_hotel":
            return self.report_summary_by_hotel(sales)
        if self.type == "sales_consecutive":
            return self.report_sales_consecutive(sales)
        if self.type == "summary_by_representatives_services":
            return self.report_summary_by_representatives_services(sales)
        if self.type == "summary_by_sale_types":
            return self.report_summary_by_sale_types(sales)
        if self.type == "report_sale_cost_daily":
            return self.report_sale_cost_daily(sales)
        if self.type == "report_by_payment_method":
            return self.report_by_payment_method(sales)
        if self.type == "summary_by_providers_services":
            return self.report_summary_by_providers_services(sales)
        if self.type == "report_sale_by_sale_type_and_hotel":
            return self.report_sale_by_sale_type_and_hotel(sales)
        if self.type == "report_refund_by_representatives":
            return self.report_refund_by_representatives(sales)
        if self.type == "report_sales_with_discount":
            return self.report_sales_with_discount(sales)
        if self.type == "report_sales_by_representatives_and_providers":
            return self.report_sales_by_representatives_and_providers(sales)
        if self.type == "report_sales_pax_by_services":
            return self.report_sales_pax_by_services(sales)
    
    def report_summary_by_providers_services(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'total_pax':        sales.paxTotal(),
            'total_import':     0,
            'sale_types':       [],
        }
        sale_type_ids = sales.order_by('sale_type').values_list('sale_type', flat=True).distinct('sale_type')
        for sale_type_id in sale_type_ids:
            sale_type = SaleType.objects.get(id=sale_type_id)
            sales_by_sale_types = sales.filter(status__in=["A","R"],sale_type=sale_type)
            sale_type_group = {
                'name':sale_type.name,
                'total_pax':sales_by_sale_types.paxTotal(),
                'total_import':0,
                'providers': []
            }
            provider_ids = sales.order_by('service__provider').values_list('service__provider', flat=True).distinct('service__provider')
            for provider_id in provider_ids:
                provider = Provider.objects.get(id=provider_id)
                sales_by_provider = sales_by_sale_types.filter(service__provider=provider)
                provider_group = {
                    'name':provider.name,
                    'total_pax':sales_by_provider.paxTotal(),
                    'total_import':0,
                    'services': []
                }
                service_ids = sales_by_provider.order_by('service').values_list('service', flat=True).distinct('service')
                for service_id in service_ids:
                    service = Service.objects.get(id=service_id)
                    sales_by_service = sales_by_provider.filter(service=service)
                    service_group = {
                        'name':service.name,
                        'provider':"{} {}".format(service.provider.name,service.provider.business_name),
                        'total_pax':sales_by_service.paxTotal(),
                        'total_import':0,
                    }
                    for sale in sales_by_service:
                        totals = sale_subtotal(sale)
                        service_group['total_import'] += totals['total_num']
                    provider_group['services'].append(service_group)
                    provider_group['total_import'] += service_group['total_import']
                sale_type_group['providers'].append(provider_group)
                sale_type_group['total_import'] += provider_group['total_import']
            context['sale_types'].append(sale_type_group)
            context['total_import'] += sale_type_group['total_import']
            print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_summary_by_providers_services_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "VENTAS POR TIPO DE VENTA PROVEEDOR Y SERVICIO ({}): {} .pdf".format(self.date_filter,date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:K', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:F2", "VENTAS POR TIPO DE VENTA, PROVEEDOR Y SERVICIO", merge_format)
            worksheet.merge_range("D3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("G1:H2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            row = 3
            for sale_type_group in context['sale_types']:
                worksheet.merge_range(row,3,row,4, sale_type_group['name'], merge_format)
                row += 1
                worksheet.merge_range(row,0,row,1, "SERVICIO", header_format)
                worksheet.merge_range(row,2,row,4, "PROVEEDOR", header_format)
                worksheet.write(row,5, "Adultos", header_format)
                worksheet.write(row,6, "Men", header_format)
                worksheet.write(row,7, "IMPORTE", header_format)
                row += 1
                for provider_group in sale_type_group['providers']:
                    worksheet.merge_range(row,0,row,2, provider_group['name'])
                    row += 1
                    for service in provider_group['services']:
                        worksheet.merge_range(row,0,row,1, service['name'])
                        worksheet.merge_range(row,2,row,4, service['provider'])
                        worksheet.write(row,5, int(service['total_pax']['adults_total']))
                        worksheet.write(row,6, int(service['total_pax']['childs_total']))
                        worksheet.write(row,7, service['total_import'])
                        row += 1
                    worksheet.write(row,5, int(provider_group['total_pax']['adults_total']),totals_format)
                    worksheet.write(row,6, int(provider_group['total_pax']['childs_total']),totals_format)
                    worksheet.write(row,7, provider_group['total_import'],totals_format)
                    row += 1
                worksheet.write(row,5, int(sale_type_group['total_pax']['adults_total']),totals_format)
                worksheet.write(row,6, int(sale_type_group['total_pax']['childs_total']),totals_format)
                worksheet.write(row,7, sale_type_group['total_import'],totals_format)
                row += 1
            worksheet.write(row,5, int(context['total_pax']['adults_total']),totals_format)
            worksheet.write(row,6, int(context['total_pax']['childs_total']),totals_format)
            worksheet.write(row,7, context['total_import'],totals_format)
            workbook.close()
            output.seek(0)
            filename = "VENTAS POR TIPO DE VENTA PROVEEDOR Y SERVICIO ({}): {} .xlsx".format(self.date_filter,date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_detail_by_representative(self,sales):
        representatives = sales.order_by('representative').values_list('representative', flat=True).distinct('representative')
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'with_out_tax':     self.with_out_tax,
            'representatives': [],
        }
        for representative in representatives:
            sales_by_representative = sales.filter(representative__id=representative)
            representative_group = {
                'representative':RepresentativeSerializer(Representative.objects.get(id=representative)).data,
                'pax_total':sales_by_representative.filter(status__in=["A","R"]).paxTotal(),
                'total_import':0,
                'total_discount':0,
                'total':0,
                'total_direct_sale':0,
                'total_cr_crc':0,
                'total_no_comis':0,
                'total_comis_rep':0,
                'dates':[]
            }
            dates = sales_by_representative.saleGroupByDay()
            for date in dates:
                sales_by_date = sales_by_representative.filter(sale_date__date=date).order_by("sale_date")
                date_sale_group = {
                    'date':date,
                    'pax_total':sales_by_date.filter(status__in=["A","R"]).paxTotal(),
                    'total_import':0,
                    'total_discount':0,
                    'total':0,
                    'total_direct_sale':0,
                    'total_cr_crc':0,
                    'total_no_comis':0,
                    'total_comis_rep':0,
                    'sales':[]
                }
                for sale in sales_by_date:
                    sale_data = sale_subtotal(sale,self.with_out_tax)
                    date_sale_group['total_import'] += sale_data['subtotal_num']
                    date_sale_group['total_discount'] += sale_data['discount_num']
                    date_sale_group['total'] += sale_data['total_num']
                    date_sale_group['total_direct_sale'] += sale_data['direct_sale']
                    date_sale_group['total_cr_crc'] += sale_data['cr_crc']
                    date_sale_group['total_no_comis'] += sale_data['no_comis']
                    date_sale_group['total_comis_rep'] += sale_data['rep_comission']
                    date_sale_group['sales'].append({
                        'sale':sale,
                        'totals':sale_data
                    })
                representative_group['total_import'] += date_sale_group['total_import']
                representative_group['total_discount'] += date_sale_group['total_discount']
                representative_group['total'] += date_sale_group['total']
                representative_group['total_direct_sale'] += date_sale_group['total_direct_sale']
                representative_group['total_cr_crc'] += date_sale_group['total_cr_crc']
                representative_group['total_no_comis'] += date_sale_group['total_no_comis']
                representative_group['total_comis_rep'] += date_sale_group['total_comis_rep']
                representative_group['dates'].append(date_sale_group)
            context['representatives'].append(representative_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('sale_report_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "Detalles de representante: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            total_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
            })
            worksheet = workbook.add_worksheet()
            worksheet.merge_range("A1:C1", self.property.name, merge_format)
            worksheet.merge_range("D1:H1", "DETALLES DE VENTAS Y COMISIONES POR REP. {}".format(" - SIN IVA " if self.with_out_tax is True else ""), merge_format)
            worksheet.merge_range("D2:H2", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.write("I1", self.request.user.username, merge_format)
            worksheet.write("I2", date_time_now.strftime("%d/%m/%Y %H:%M:%S"), merge_format)
            worksheet.set_column('A:H', 20)
            worksheet.write("B3", "REPRESENTANTE", header_format)
            worksheet.write("C3", "ACTIVIDAD", header_format)
            worksheet.write("D3", "GRUPO", header_format)
            worksheet.write("E3", "CUPON", header_format)
            worksheet.write("F3", "SERVICIO", header_format)
            worksheet.set_column('F:F', 50)
            worksheet.write("G3", "PAX", header_format)
            worksheet.set_column('G:G', 10)
            worksheet.write("H3", "Fch Serv", header_format)
            worksheet.set_column('I:Q', 15)
            worksheet.write("I3", "IMPORTE", header_format)
            worksheet.write("J3", "%", header_format)
            worksheet.write("K3", "Descuento", header_format)
            worksheet.write("L3", "TOTAL", header_format)
            worksheet.write("M3", "VTA.DIRECTA", header_format)
            worksheet.write("N3", "CR/CRC", header_format)
            worksheet.write("O3", "NO COMIS", header_format)
            worksheet.write("P3", "COMIS REP", header_format)
            worksheet.write("Q3", "HAB", header_format)
            row = 3
            for representative_group in context['representatives']:
                for date_sale_group in representative_group['dates']:
                    worksheet.write(row,0, date_sale_group['date'].strftime("%d/%m/%Y"))
                    row += 1
                    for sale_data in date_sale_group['sales']:
                        sale = sale_data['sale']
                        totals = sale_data['totals']
                        worksheet.write(row,1, "{} - {} ".format(sale.representative.code,sale.representative.name))
                        if sale.service.activity is not None:
                            worksheet.write(row,2, sale.service.activity.name)
                        if sale.service.business_group is not None:
                            worksheet.write(row,3, sale.service.business_group.name)
                        worksheet.write(row,4, "R{}".format(str(sale.sale_key).zfill(8)) if sale.status == "R" else str(sale.sale_key).zfill(8))
                        worksheet.write(row,5, "^^^ C A N C E L A D O ^^^" if sale.status == "C" else sale.service.name)
                        if sale.status != "C" :
                            worksheet.write(row,6, "{}.{}".format(int(sale.adults),int(sale.childs)))
                            worksheet.write(row,7, sale.service_date.strftime("%d/%m/%Y"))
                            worksheet.write(row,8, totals['subtotal_num'])
                            if sale.discount_type == "percent":
                                worksheet.write(row,9, sale.discount)
                            worksheet.write(row,10, totals['discount_num'])
                            worksheet.write(row,11, totals['total_num'])
                            worksheet.write(row,12, totals['direct_sale'])
                            worksheet.write(row,13, totals['cr_crc'])
                            worksheet.write(row,14, totals['no_comis'])
                            worksheet.write(row,15, totals['rep_comission'])
                            worksheet.write(row,16, sale.room)
                        row += 1
                    worksheet.write(row,6, "{}.{}".format(int(date_sale_group['pax_total']['adults_total']),int(date_sale_group['pax_total']['childs_total'])), total_format)
                    worksheet.write(row,8, date_sale_group['total_import'], total_format)
                    worksheet.write(row,10, date_sale_group['total_discount'], total_format)
                    worksheet.write(row,11, date_sale_group['total'], total_format)
                    worksheet.write(row,12, date_sale_group['total_direct_sale'], total_format)
                    worksheet.write(row,13, date_sale_group['total_cr_crc'], total_format)
                    worksheet.write(row,14, date_sale_group['total_no_comis'], total_format)
                    worksheet.write(row,15, date_sale_group['total_comis_rep'], total_format)
                    row += 1
                worksheet.merge_range(row,0,row,5, "{} - {} ".format(representative_group['representative']['code'],representative_group['representative']['name']))
                worksheet.write(row,6, "{}.{}".format(int(representative_group['pax_total']['adults_total']),int(representative_group['pax_total']['childs_total'])), total_format)
                worksheet.write(row,8, representative_group['total_import'], total_format)
                worksheet.write(row,10, representative_group['total_discount'], total_format)
                worksheet.write(row,11, representative_group['total'], total_format)
                worksheet.write(row,12, representative_group['total_direct_sale'], total_format)
                worksheet.write(row,13, representative_group['total_cr_crc'], total_format)
                worksheet.write(row,14, representative_group['total_no_comis'], total_format)
                worksheet.write(row,15, representative_group['total_comis_rep'], total_format)
                row += 1
            workbook.close()
            output.seek(0)
            filename = 'Detalles de representante: {} .xlsx'.format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
                            
    def report_summary_by_representative(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'representatives': [],
        }
        representatives = sales.order_by('representative').values_list('representative', flat=True).distinct('representative')
        for representative in representatives:
            sales_by_representative = sales.filter(status__in=["A","R"],representative__id=representative)
            representative_group = {
                'representative':RepresentativeSerializer(Representative.objects.get(id=representative)).data,
                'total_import_without_tax':0,
                'total_import':0,
                'total_discount':0,
                'total_comission':0,
            }
            for sale in sales_by_representative:
                totals = sale_subtotal(sale)
                totals_notax = sale_subtotal(sale,True)
                representative_group['total_import'] += totals['total_num']
                representative_group['total_import_without_tax'] += totals_notax['total_num']
                representative_group['total_discount'] += totals['discount_num']
                representative_group['total_comission'] += totals['rep_comission']
            context['representatives'].append(representative_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('sale_summary_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "Resumen de representantes: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'top':1,
                'bottom':1,
                'right':1,
                'left':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            header_format_1 = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:F2", "RESUMEN DE VENTAS Y COMISIONES", merge_format)
            worksheet.merge_range("D3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("G1:H2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            worksheet.merge_range("E4:F4", "VENTAS NETAS", header_format)
            worksheet.merge_range("A5:D5", "REPRESENTANTE", header_format_1)
            worksheet.write("E5", "CON IVA", header_format)
            worksheet.write("F5", "SIN IVA", header_format)
            worksheet.write("G5", "DESCTOS", header_format)
            worksheet.write("H5", "COMISIÓN", header_format)
            row = 5
            for representative_group in context['representatives']:
                worksheet.merge_range(row,0,row,3, "{} - {} ".format(representative_group['representative']['code'],representative_group['representative']['name']))
                worksheet.write(row,4, representative_group['total_import'])
                worksheet.write(row,5, representative_group['total_import_without_tax'])
                worksheet.write(row,6, representative_group['total_discount'])
                worksheet.write(row,7, representative_group['total_comission'])
                row += 1
            workbook.close()
            output.seek(0)
            filename = 'Resumen de representantes: {} .xlsx'.format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_summary_by_sales_types_services(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'date_filter':      self.date_filter,
            'just_import':      self.just_import,
            'total_import':0,
            'total_cost':0,
            'total_comission':0,
            'total_comission_sup':0,
            'total_comission_coord':0,
            'total_utility':0,
            'sale_types': [],
        }
        sale_type_ids = sales.filter(status__in=["A","R"]).order_by('sale_type').values_list('sale_type', flat=True).distinct('sale_type')
        for sale_type_id in sale_type_ids:
            sale_type = SaleType.objects.get(id=sale_type_id)
            sales_by_sale_types = sales.filter(sale_type=sale_type)
            sale_type_group = {
                'name':sale_type.name,
                'total_pax':sales_by_sale_types.paxTotal(),
                'total_import':0,
                'total_cost':0,
                'total_comission':0,
                'total_comission_sup':0,
                'total_comission_coord':0,
                'total_utility':0,
                'services': []
            }
            service_ids = sales_by_sale_types.order_by('service').values_list('service', flat=True).distinct('service')
            for service_id in service_ids:
                service = Service.objects.get(id=service_id)
                sales_by_service = sales_by_sale_types.filter(service=service)
                service_group = {
                    'name':service.name,
                    'activity':service.activity.name if service.activity is not None else "",
                    'business_group':service.business_group.name if service.business_group is not None else "",
                    'total_pax':sales_by_service.paxTotal(),
                    'total_import':0,
                    'total_cost':0,
                    'total_comission':0,
                    'total_comission_sup':0,
                    'total_comission_coord':0,
                    'total_utility':0,
                }
                for sale in sales_by_service:
                    totals = sale_subtotal(sale)
                    #if totals['total_cost_num'] != 0:
                    service_group['total_import'] += totals['total_num']
                    service_group['total_cost'] += totals['total_cost_num']
                    service_group['total_comission'] += totals['rep_comission']
                    service_group['total_comission_sup'] += totals['sup_comis']
                    service_group['total_comission_coord'] += totals['coords_comis']
                    utility = totals['total_num'] - (totals['total_cost_num']+totals['rep_comission']+totals['sup_comis']+totals['coords_comis']) 
                    service_group['total_utility'] += utility
                sale_type_group['services'].append(service_group)
                sale_type_group['total_import'] += service_group['total_import']
                sale_type_group['total_cost'] += service_group['total_cost']
                sale_type_group['total_comission'] += service_group['total_comission']
                sale_type_group['total_comission_sup'] += service_group['total_comission_sup']
                sale_type_group['total_comission_coord'] += service_group['total_comission_coord']
                sale_type_group['total_utility'] += service_group['total_utility']
            context['sale_types'].append(sale_type_group)
            context['total_import'] += sale_type_group['total_import']
            context['total_cost'] += sale_type_group['total_cost']
            context['total_comission'] += sale_type_group['total_comission']
            context['total_comission_sup'] += sale_type_group['total_comission_sup']
            context['total_comission_coord'] += sale_type_group['total_comission_coord']
            context['total_utility'] += sale_type_group['total_utility']
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_summary_by_sales_types_services_template.html')
            
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "Ventas por agencia y servicios ({}): {} .pdf".format(self.date_filter,date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:K', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:F2", "VENTAS POR TIPO DE VENTA Y SERVICIOS ({})".format(self.date_filter), merge_format)
            worksheet.merge_range("D3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("G1:H2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            row = 3
            for sale_type in context['sale_types']:
                worksheet.merge_range(row,3,row,4, sale_type['name'], merge_format)
                row += 1
                worksheet.merge_range(row,0,row,1, "SERVICIO", header_format)
                worksheet.write(row,2, "ACTIVIDAD", header_format)
                worksheet.write(row,3, "GRUPO", header_format)
                worksheet.write(row,4, "PAX", header_format)
                worksheet.write(row,5, "IMPORTE", header_format)
                if self.just_import is False:
                    worksheet.write(row,6, "COSTO", header_format)
                    worksheet.write(row,7, "COMIS REP", header_format)
                    worksheet.write(row,8, "COM. SUP", header_format)
                    worksheet.write(row,9, "COMIS. COORD.", header_format)
                    worksheet.write(row,10, "UTILIDAD", header_format)
                row += 1
                for service in sale_type['services']:
                    worksheet.merge_range(row,0,row,1, service['name'])
                    worksheet.write(row,2, service['activity'])
                    worksheet.write(row,3, service['business_group'])
                    worksheet.write(row,4, "{}.{}".format(int(service['total_pax']['adults_total']),int(service['total_pax']['childs_total'])))
                    worksheet.write(row,5, service['total_import'])
                    if self.just_import is False:
                        worksheet.write(row,6, service['total_cost'])
                        worksheet.write(row,7, service['total_comission'])
                        worksheet.write(row,8, service['total_comission_sup'])
                        worksheet.write(row,9, service['total_comission_coord'])
                        worksheet.write(row,10, service['total_utility'])
                    row += 1
                worksheet.write(row,4, "{}.{}".format(int(sale_type['total_pax']['adults_total']),int(sale_type['total_pax']['childs_total'])),totals_format)
                worksheet.write(row,5, sale_type['total_import'],totals_format)
                if self.just_import is False:
                    worksheet.write(row,6, sale_type['total_cost'],totals_format)
                    worksheet.write(row,7, sale_type['total_comission'],totals_format)
                    worksheet.write(row,8, sale_type['total_comission_sup'],totals_format)
                    worksheet.write(row,9, sale_type['total_comission_coord'],totals_format)
                    worksheet.write(row,10, sale_type['total_utility'],totals_format)
                row += 1
            worksheet.merge_range(row,0,row,1, "TOTAL", header_format)
            worksheet.write(row,5, context['total_import'],totals_format)
            if self.just_import is False:
                worksheet.write(row,6, context['total_cost'],totals_format)
                worksheet.write(row,7, context['total_comission'],totals_format)
                worksheet.write(row,8, context['total_comission_sup'],totals_format)
                worksheet.write(row,9, context['total_comission_coord'],totals_format)
                worksheet.write(row,10, context['total_utility'],totals_format)
            workbook.close()
            output.seek(0)
            filename = "Ventas por agencia y servicios ({}): {} .xlsx".format(self.date_filter,date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_summary_by_services(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'date_filter':      self.date_filter,
            'just_import':      self.just_import,
            'providers': [],
        }
        provider_ids = sales.order_by('service__provider').values_list('service__provider', flat=True).distinct('service__provider')
        for provider_id in provider_ids:
            provider = Provider.objects.get(id=provider_id)
            sales_by_provider = sales.filter(status__in=["A","R"],service__provider=provider)
            provider_group = {
                'name':provider.name,
                'business_name':provider.business_name,
                'total_pax':sales_by_provider.paxTotal(),
                'total_import':0,
                'total_comission_hc':0,
                'total_cost':0,
                'total_comission':0,
                'services': []
            }
            service_ids = sales_by_provider.order_by('service').values_list('service', flat=True).distinct('service')
            for service_id in service_ids:
                service = Service.objects.get(id=service_id)
                sales_by_service = sales_by_provider.filter(service=service)
                service_group = {
                    'name':service.name,
                    'total_pax':sales_by_service.paxTotal(),
                    'total_import':0,
                    'total_cost':0,
                    'total_comission':0,
                    'total_comission_hc':0,
                }
                for sale in sales_by_service:
                    totals = sale_subtotal(sale)
                    service_group['total_import'] += totals['total_num']
                    service_group['total_cost'] += totals['total_cost_num']
                    service_group['total_comission'] += totals['rep_comission']
                    comission_hc = totals['total_num'] - totals['total_cost_num']
                    service_group['total_comission_hc'] += comission_hc
                provider_group['services'].append(service_group)
                provider_group['total_import'] += service_group['total_import']
                provider_group['total_cost'] += service_group['total_cost']
                provider_group['total_comission'] += service_group['total_comission']
                provider_group['total_comission_hc'] += service_group['total_comission_hc']
            context['providers'].append(provider_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_summary_by_services_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "Resumen por servicio ({}): {} .pdf".format(self.date_filter,date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            provider_format = workbook.add_format({
                'bold': True,
                "align": "right",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:J', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:F2", "RESUMEN POR SERVICIO ({})".format(self.date_filter), merge_format)
            worksheet.merge_range("D3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("G1:H2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            worksheet.merge_range("A4:B4", "PROVEEDOR", header_format)
            worksheet.merge_range("C4:D4", "SERVICIO", header_format)
            worksheet.write("E4", "Adultos", header_format)
            worksheet.write("F4", "Men", header_format)
            worksheet.write("G4", "IMPORTE", header_format)
            if self.just_import is False:
                worksheet.write("H4", "COMISION AG", header_format)
                worksheet.write("I4", "COSTO", header_format)
                worksheet.write("J4", "COMIS. REP.", header_format)
            row = 4
            for provider in context['providers']:
                for service in provider['services']:
                    worksheet.merge_range(row,0,row,1, "{} {}".format(provider['name'],provider['business_name']))
                    worksheet.merge_range(row,2,row,3, service['name'])
                    worksheet.write(row,4, int(service['total_pax']['adults_total']))
                    worksheet.write(row,5, int(service['total_pax']['childs_total']))
                    worksheet.write(row,6, service['total_import'])
                    if self.just_import is False:
                        worksheet.write(row,7, service['total_comission_hc'])
                        worksheet.write(row,8, service['total_cost'])
                        worksheet.write(row,9, service['total_comission'])
                    row += 1
                worksheet.merge_range(row,0,row,3, "{} {}".format(provider['name'],provider['business_name']), provider_format)
                worksheet.write(row,4, int(provider['total_pax']['adults_total']),totals_format)
                worksheet.write(row,5, int(provider['total_pax']['childs_total']),totals_format)
                worksheet.write(row,6, provider['total_import'],totals_format)
                if self.just_import is False:
                    worksheet.write(row,7, provider['total_comission_hc'],totals_format)
                    worksheet.write(row,8, provider['total_cost'],totals_format)
                    worksheet.write(row,9, provider['total_comission'],totals_format)
                row += 1
            workbook.close()
            output.seek(0)
            filename = "Resumen por servicio ({}): {} .xlsx".format(self.date_filter,date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_summary_by_hotel(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'hotels': [],
            'total_pax':sales.filter(status__in=["A","R"]).paxTotal(),
            'total_import':0,
        }
        hotel_ids = sales.order_by('hotel').values_list('hotel', flat=True).distinct('hotel')
        for hotel_id in hotel_ids:
            hotel = Hotel.objects.get(id=hotel_id)
            sales_by_hotel = sales.filter(status__in=["A","R"],hotel=hotel)
            hotel_group = {
                'name':hotel.name,
                'total_pax':sales_by_hotel.paxTotal(),
                'total_import':0,
            }
            for sale in sales_by_hotel:
                totals = sale_subtotal(sale)
                hotel_group['total_import'] += totals['total_num']
            context['total_import'] += hotel_group['total_import']
            context['hotels'].append(hotel_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_summary_by_hotel_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "VENTAS POR HOTEL: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            provider_format = workbook.add_format({
                'bold': True,
                "align": "right",
                "valign": "vcenter",
                'font_size':12,
            })
            total_format = workbook.add_format({
                'top':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:E2", "VENTAS POR HOTEL", merge_format)
            worksheet.merge_range("C3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("F1:G2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            worksheet.merge_range("A4:E4", "HOTEL", header_format)
            worksheet.write("F4", "PAX", header_format)
            worksheet.write("G4", "IMPORTE", header_format)
            row = 4
            for hotel in context['hotels']:
                worksheet.merge_range(row,0,row,4, hotel['name'])
                worksheet.write(row,5, "{}.{}".format(int(hotel['total_pax']['adults_total']),int(hotel['total_pax']['childs_total'])))
                worksheet.write(row,6, hotel['total_import'])
                row += 1
            worksheet.merge_range(row,0,row,4, "GRAN TOTAL", provider_format)
            worksheet.write(row,5, "{}.{}".format(int(hotel['total_pax']['adults_total']),int(hotel['total_pax']['childs_total'])), total_format)
            worksheet.write(row,6, hotel['total_import'], total_format)
            workbook.close()
            output.seek(0)
            filename = "VENTAS POR HOTEL: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_sales_consecutive(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'sales': [],
            'total_pax':sales.paxTotal(),
            'total_import':0,
        }
        sales = sales.order_by('sale_key')
        for sale in sales:
            totals = sale_subtotal(sale)
            sale_group = {
                'sale':sale,
                'import':totals['total_num']
            }
            context['total_import'] += totals['total_num']
            context['sales'].append(sale_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_sales_consecutive_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "CONSECUTIVO DE CUPONES: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:E2", "CONSECUTIVO DE CUPONES", merge_format)
            worksheet.merge_range("C3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("F1:G2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)

            worksheet.write("A4", "FECHA", header_format)
            worksheet.write("B4", "CUPON", header_format)
            worksheet.merge_range("C4:E4", "SERVICIO", header_format)
            worksheet.write("F4", "PAX", header_format)
            worksheet.write("G4", "IMPORTE", header_format)
            row = 4
            for sale_group in context['sales']:
                sale = sale_group['sale']
                worksheet.write(row,0, sale.sale_date.strftime("%d/%m/%Y"))
                worksheet.write(row,1, "R{}".format(str(sale.sale_key).zfill(8)) if sale.status == "R" else str(sale.sale_key).zfill(8))
                worksheet.merge_range(row,2,row,4, "^^^ C A N C E L A D O ^^^" if sale.status == "C" else sale.service.name)
                if sale.status != "C" :
                    worksheet.write(row,5, "{}.{}".format(int(sale.adults),int(sale.childs)))
                    worksheet.write(row,6, sale_group['import'])
                row += 1
            worksheet.write(row,5, "{}.{}".format(int(context['total_pax']['adults_total']),int(context['total_pax']['childs_total'])),totals_format)
            worksheet.write(row,6, context['total_import'],totals_format)
            workbook.close()
            output.seek(0)
            filename = "CONSECUTIVO DE CUPONES: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_summary_by_representatives_services(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'sale_types': [],
        }
        sale_type_ids = sales.order_by('sale_type').values_list('sale_type', flat=True).distinct('sale_type')
        for sale_type_id in sale_type_ids:
            sale_type = SaleType.objects.get(id=sale_type_id)
            sales_by_sale_types = sales.filter(status__in=["A","R"],sale_type=sale_type)
            sale_type_group = {
                'name':sale_type.name,
                'total_pax':sales_by_sale_types.paxTotal(),
                'total_import':0,
                'total_comission':0,
                'representatives': []
            }
            representative_ids = sales.order_by('representative').values_list('representative', flat=True).distinct('representative')
            for representative_id in representative_ids:
                representative = Representative.objects.get(id=representative_id)
                sales_by_representative = sales_by_sale_types.filter(representative=representative)
                representative_group = {
                    'name':representative.name,
                    'total_pax':sales_by_representative.paxTotal(),
                    'total_import':0,
                    'total_comission':0,
                    'services': []
                }
                service_ids = sales_by_representative.order_by('service').values_list('service', flat=True).distinct('service')
                for service_id in service_ids:
                    service = Service.objects.get(id=service_id)
                    sales_by_service = sales_by_representative.filter(service=service)
                    service_group = {
                        'name':service.name,
                        'provider':"{} {}".format(service.provider.name,service.provider.business_name),
                        'total_pax':sales_by_service.paxTotal(),
                        'total_import':0,
                        'total_comission':0,
                    }
                    for sale in sales_by_service:
                        totals = sale_subtotal(sale)
                        service_group['total_import'] += totals['total_num']
                        service_group['total_comission'] += totals['rep_comission']
                    representative_group['services'].append(service_group)
                    representative_group['total_import'] += service_group['total_import']
                    representative_group['total_comission'] += service_group['total_comission']
                sale_type_group['representatives'].append(representative_group)
                sale_type_group['total_import'] += representative_group['total_import']
                sale_type_group['total_comission'] += representative_group['total_comission']
            context['sale_types'].append(sale_type_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_summary_by_representatives_services_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "Resumen por servicio ({}): {} .pdf".format(self.date_filter,date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:K', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:F2", "VENTAS POR REPRESENTANTE Y SERVICIOS", merge_format)
            worksheet.merge_range("D3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("G1:H2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            row = 3
            for sale_type_group in context['sale_types']:
                worksheet.merge_range(row,3,row,4, sale_type_group['name'], merge_format)
                row += 1
                worksheet.merge_range(row,0,row,1, "SERVICIO", header_format)
                worksheet.merge_range(row,2,row,4, "PROVEEDOR", header_format)
                worksheet.write(row,5, "PAX", header_format)
                worksheet.write(row,6, "IMPORTE", header_format)
                worksheet.write(row,7, "COMISION", header_format)
                row += 1
                for representative_group in sale_type_group['representatives']:
                    worksheet.merge_range(row,0,row,2, representative_group['name'])
                    row += 1
                    for service in representative_group['services']:
                        worksheet.merge_range(row,0,row,1, service['name'])
                        worksheet.merge_range(row,2,row,4, service['provider'])
                        worksheet.write(row,5, "{}.{}".format(int(service['total_pax']['adults_total']),int(service['total_pax']['childs_total'])))
                        worksheet.write(row,6, service['total_import'])
                        worksheet.write(row,7, service['total_comission'])
                        row += 1
                    worksheet.write(row,5, "{}.{}".format(int(representative_group['total_pax']['adults_total']),int(representative_group['total_pax']['childs_total'])),totals_format)
                    worksheet.write(row,6, representative_group['total_import'],totals_format)
                    worksheet.write(row,7, representative_group['total_comission'],totals_format)
                    row += 1
                worksheet.write(row,5, "{}.{}".format(int(sale_type_group['total_pax']['adults_total']),int(sale_type_group['total_pax']['childs_total'])),totals_format)
                worksheet.write(row,6, sale_type_group['total_import'],totals_format)
                worksheet.write(row,7, sale_type_group['total_comission'],totals_format)
                row += 1
            workbook.close()
            output.seek(0)
            filename = "Resumen por servicio ({}): {} .xlsx".format(self.date_filter,date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response

    def report_summary_by_sale_types(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'sale_types': [],
            'total_pax':sales.filter(status__in=["A","R"]).paxTotal(),
            'total_import':0,
        }
        sale_type_ids = sales.order_by('sale_type').values_list('sale_type', flat=True).distinct('sale_type')
        for sale_type_id in sale_type_ids:
            sale_type = SaleType.objects.get(id=sale_type_id)
            sales_by_sale_types = sales.filter(status__in=["A","R"],sale_type=sale_type)
            sale_type_group = {
                'name':sale_type.name,
                'total_pax':sales_by_sale_types.paxTotal(),
                'total_import':0,
            }
            for sale in sales_by_sale_types:
                totals = sale_subtotal(sale)
                sale_type_group['total_import'] += totals['total_num']
            context['total_import'] += sale_type_group['total_import']
            context['sale_types'].append(sale_type_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_summary_by_sale_types_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "VENTAS POR TIPO DE VENTA: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            provider_format = workbook.add_format({
                'bold': True,
                "align": "right",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:K', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:F2", "VENTAS POR TIPO DE VENTA", merge_format)
            worksheet.merge_range("D3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("G1:H2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            worksheet.merge_range("B4:E4", "TIPO DE VENTA", header_format)
            worksheet.write("F4", "PAX", header_format)
            worksheet.write("G4", "IMPORTE", header_format)
            row = 4
            for sale_type_group in context['sale_types']:
                worksheet.merge_range(row,1,row,4, sale_type_group['name'])
                worksheet.write(row,5, "{}.{}".format(int(sale_type_group['total_pax']['adults_total']),int(sale_type_group['total_pax']['childs_total'])))
                worksheet.write(row,6, sale_type_group['total_import'])
                row += 1
            worksheet.merge_range(row,1,row,4, "GRAN TOTAL", provider_format)
            worksheet.write(row,5, "{}.{}".format(int(context['total_pax']['adults_total']),int(context['total_pax']['childs_total'])),totals_format)
            worksheet.write(row,6, context['total_import'],totals_format)                
            workbook.close()
            output.seek(0)
            filename = "VENTAS POR TIPO DE VENTA: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
    
    def report_sale_cost_daily(self,sales):
        sales = sales.filter(status__in=["A","R"])
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'summary':          self.summary,
            'sale_types': [],
            'adults':0,
            'childs':0,
            'subtotal_num':0,
            'subtotal_num_no_tax':0,
            'discount_num':0,
            'total_num':0,
            'discount_num_no_tax':0,
            'adult_fee_num':0,
            'child_fee_num':0,
            'total_cost_num':0,
            'total_cost_num_no_tax':0,
            'direct_sale':0,
            'cr_crc':0,
            'no_comis':0,
        }
        sale_type_ids = sales.order_by('sale_type').values_list('sale_type', flat=True).distinct('sale_type')
        for sale_type_id in sale_type_ids:
            sale_type = SaleType.objects.get(id=sale_type_id)
            sales_by_sale_types = sales.filter(sale_type=sale_type).order_by('-sale_key')
            sale_type_group = {
                'name':sale_type.name,
                'dates':[],
                'adults':0,
                'childs':0,
                'subtotal_num':0,
                'subtotal_num_no_tax':0,
                'discount_num':0,
                'total_num':0,
                'discount_num_no_tax':0,
                'adult_fee_num':0,
                'child_fee_num':0,
                'total_cost_num':0,
                'total_cost_num_no_tax':0,
                'direct_sale':0,
                'cr_crc':0,
                'no_comis':0,
            }
            dates = sales_by_sale_types.saleGroupByDay()
            for date in dates:
                sales_by_date = sales_by_sale_types.filter(sale_date__date=date).order_by("sale_date")
                total_pax = sales_by_date.paxTotal()
                date_group = {
                    'sales':[],
                    'adults':total_pax['adults_total'],
                    'childs':total_pax['childs_total'],
                    'subtotal_num':0,
                    'subtotal_num_no_tax':0,
                    'discount_num':0,
                    'total_num':0,
                    'discount_num_no_tax':0,
                    'adult_fee_num':0,
                    'child_fee_num':0,
                    'total_cost_num':0,
                    'total_cost_num_no_tax':0,
                    'direct_sale':0,
                    'cr_crc':0,
                    'no_comis':0,
                }
                for sale in sales_by_date:
                    sale_group = {
                        'sale':sale,
                    }
                    sale_group.update(sale_subtotal(sale,self.with_out_tax))
                    sale_group['subtotal_num_no_tax'] = sale_group['subtotal_num']/((sale.service_rate.tax/100)+1)
                    sale_group['discount_num_no_tax'] = sale_group['discount_num']/((sale.service_rate.tax/100)+1)
                    sale_group['adult_cost_percent'] = 100 - sale.service_rate.hard_rock_comission_adult
                    sale_group['child_cost_percent'] = 100 - sale.service_rate.hard_rock_comission_child
                    date_group['sales'].append(sale_group)
                    date_group['subtotal_num'] += sale_group['subtotal_num']
                    date_group['subtotal_num_no_tax'] += sale_group['subtotal_num_no_tax']
                    date_group['discount_num'] += sale_group['discount_num']
                    date_group['total_num'] += sale_group['total_num']
                    date_group['discount_num_no_tax'] += sale_group['discount_num_no_tax']
                    date_group['adult_fee_num'] += sale_group['adult_fee_num']
                    date_group['child_fee_num'] += sale_group['child_fee_num']
                    date_group['total_cost_num'] += sale_group['total_cost_num']
                    date_group['total_cost_num_no_tax'] += sale_group['total_cost_num_no_tax']
                    date_group['direct_sale'] += sale_group['direct_sale_cost']
                    date_group['cr_crc'] += sale_group['cr_crc_cost']
                    date_group['no_comis'] += sale_group['no_comis_cost']
                sale_type_group['dates'].append(date_group)
                sale_type_group['adults'] += date_group['adults']
                sale_type_group['childs'] += date_group['childs']
                sale_type_group['subtotal_num'] += date_group['subtotal_num']
                sale_type_group['subtotal_num_no_tax'] += date_group['subtotal_num_no_tax']
                sale_type_group['discount_num'] += date_group['discount_num']
                sale_type_group['total_num'] += date_group['total_num']
                sale_type_group['discount_num_no_tax'] += date_group['discount_num_no_tax']
                sale_type_group['adult_fee_num'] += date_group['adult_fee_num']
                sale_type_group['child_fee_num'] += date_group['child_fee_num']
                sale_type_group['total_cost_num'] += date_group['total_cost_num']
                sale_type_group['total_cost_num_no_tax'] += date_group['total_cost_num_no_tax']
                sale_type_group['direct_sale'] += date_group['direct_sale']
                sale_type_group['cr_crc'] += date_group['cr_crc']
                sale_type_group['no_comis'] += date_group['no_comis']
            context['sale_types'].append(sale_type_group)
            context['adults'] += sale_type_group['adults']
            context['childs'] += sale_type_group['childs']
            context['subtotal_num'] += sale_type_group['subtotal_num']
            context['subtotal_num_no_tax'] += sale_type_group['subtotal_num_no_tax']
            context['discount_num'] += sale_type_group['discount_num']
            context['total_num'] += sale_type_group['total_num']
            context['discount_num_no_tax'] += sale_type_group['discount_num_no_tax']
            context['adult_fee_num'] += sale_type_group['adult_fee_num']
            context['child_fee_num'] += sale_type_group['child_fee_num']
            context['total_cost_num'] += sale_type_group['total_cost_num']
            context['total_cost_num_no_tax'] += sale_type_group['total_cost_num_no_tax']
            context['direct_sale'] += sale_type_group['direct_sale']
            context['cr_crc'] += sale_type_group['cr_crc']
            context['no_comis'] += sale_type_group['no_comis']
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_sale_cost_daily_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "Detalle Venta y Costo Diario {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:K', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:F2", "Detalle Venta y Costo Diario.{}".format(" - SIN IVA "if self.with_out_tax is True else ""), merge_format)
            worksheet.merge_range("D3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("G1:H2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            row = 3
            for sale_type_group in context['sale_types']:
                worksheet.merge_range(row,3,row,4, sale_type_group['name'], merge_format)
                row += 1
                worksheet.write(row,0, "CUPON", header_format)
                worksheet.write(row,1, "Nombre Pax", header_format)
                worksheet.write(row,2, "Clave", header_format)
                worksheet.write(row,3, "Servicio", header_format)
                worksheet.write(row,4, "Grupo", header_format)
                worksheet.write(row,5, "Fch Venta", header_format)
                worksheet.write(row,6, "Ad", header_format)
                worksheet.write(row,7, "Men", header_format)
                worksheet.write(row,8, "Nombre Rep", header_format)
                worksheet.write(row,9, "Venta Neta", header_format)
                worksheet.write(row,10, "sin IVA", header_format)
                worksheet.write(row,11, "Descuento", header_format)
                worksheet.write(row,12, "%", header_format)
                worksheet.write(row,13, "Venta Bruta", header_format)
                worksheet.write(row,14, "IVA Desct", header_format)
                worksheet.write(row,15, "%cost ad", header_format)
                worksheet.write(row,16, "%cost men", header_format)
                worksheet.write(row,17, "Costo Ad", header_format)
                worksheet.write(row,18, "Costo Men", header_format)
                worksheet.write(row,19, "Costo total con IVA", header_format)
                worksheet.write(row,20, "Costo total sin IVA", header_format)
                worksheet.write(row,21, "Directos sin IVA", header_format)
                worksheet.write(row,22, "RC sin IVA", header_format)
                worksheet.write(row,23, "No Comis sin IVA", header_format)
                row += 1
                for date_group in sale_type_group['dates']:
                    for sale_group in date_group['sales']:
                        sale = sale_group['sale']
                        worksheet.write(row,0, "R{}".format(str(sale.sale_key).zfill(8)) if sale.status == "R" else str(sale.sale_key).zfill(8))
                        worksheet.write(row,1, sale.name_pax)
                        worksheet.write(row,2, str(sale.service.id).zfill(4))
                        worksheet.write(row,3, sale.service.name)
                        if sale.service.business_group is not None:
                            worksheet.write(row,4, sale.service.business_group.name)
                        worksheet.write(row,5, sale.sale_date.strftime("%d/%m/%Y"))
                        worksheet.write(row,6, int(sale.adults))
                        worksheet.write(row,7, int(sale.childs))
                        worksheet.write(row,8, sale.representative.name)
                        worksheet.write(row,9, sale_group['subtotal_num'])
                        worksheet.write(row,10, sale_group['subtotal_num_no_tax'])
                        worksheet.write(row,11, sale_group['discount_num'])
                        worksheet.write(row,12, sale.discount)
                        worksheet.write(row,13, sale_group['total_num'])
                        worksheet.write(row,14, sale_group['discount_num_no_tax'])
                        worksheet.write(row,15, sale_group['adult_cost_percent'])
                        worksheet.write(row,16, sale_group['child_cost_percent'])
                        worksheet.write(row,17, sale_group['adult_fee_num'])
                        worksheet.write(row,18, sale_group['child_fee_num'])
                        worksheet.write(row,19, sale_group['total_cost_num'])
                        worksheet.write(row,20, sale_group['total_cost_num_no_tax'])
                        worksheet.write(row,21, sale_group['direct_sale_cost'])
                        worksheet.write(row,22, sale_group['cr_crc_cost'])
                        worksheet.write(row,23, sale_group['no_comis_cost'])
                        row += 1
                    worksheet.write(row,6, int(date_group['adults']),totals_format)
                    worksheet.write(row,7, int(date_group['childs']),totals_format)
                    worksheet.write(row,9, date_group['subtotal_num'],totals_format)
                    worksheet.write(row,10, date_group['subtotal_num_no_tax'],totals_format)
                    worksheet.write(row,11, date_group['discount_num'],totals_format)
                    worksheet.write(row,13, date_group['total_num'],totals_format)
                    worksheet.write(row,14, date_group['discount_num_no_tax'],totals_format)
                    worksheet.write(row,17, date_group['adult_fee_num'],totals_format)
                    worksheet.write(row,18, date_group['child_fee_num'],totals_format)
                    worksheet.write(row,19, date_group['total_cost_num'],totals_format)
                    worksheet.write(row,20, date_group['total_cost_num_no_tax'],totals_format)
                    worksheet.write(row,21, date_group['direct_sale'],totals_format)
                    worksheet.write(row,22, date_group['cr_crc'],totals_format)
                    worksheet.write(row,23, date_group['no_comis'],totals_format)
                    row += 1
                worksheet.write(row,6, int(sale_type_group['adults']),totals_format)
                worksheet.write(row,7, int(sale_type_group['childs']),totals_format)
                worksheet.write(row,9, sale_type_group['subtotal_num'],totals_format)
                worksheet.write(row,10, sale_type_group['subtotal_num_no_tax'],totals_format)
                worksheet.write(row,11, sale_type_group['discount_num'],totals_format)
                worksheet.write(row,13, sale_type_group['total_num'],totals_format)
                worksheet.write(row,14, sale_type_group['discount_num_no_tax'],totals_format)
                worksheet.write(row,17, sale_type_group['adult_fee_num'],totals_format)
                worksheet.write(row,18, sale_type_group['child_fee_num'],totals_format)
                worksheet.write(row,19, sale_type_group['total_cost_num'],totals_format)
                worksheet.write(row,20, sale_type_group['total_cost_num_no_tax'],totals_format)
                worksheet.write(row,21, sale_type_group['direct_sale'],totals_format)
                worksheet.write(row,22, sale_type_group['cr_crc'],totals_format)
                worksheet.write(row,23, sale_type_group['no_comis'],totals_format)
                row += 1
            workbook.close()
            output.seek(0)
            filename = 'Detalle Venta y Costo Diario {} .xlsx'.format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
                       
    def report_by_payment_method(self,sales):
        sales = sales.filter(status__in=["A","R","C"])
        representatives = sales.order_by('representative').values('representative__name','representative__id','representative__code').distinct('representative')
        payment_methods = models.PaymentMethod.objects.filter(sale_payments__sale__in=sales).values('id','name').distinct()
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'summary':          self.summary,
            'payment_methods':  payment_methods,
            'total':{
                'total_import':     0,
                'total_import_mn':  0,
            },
            'representatives':  [],
        }
        for payment_method in payment_methods:
            context["total"]["total_"+str(payment_method['id'])] = 0
        for representative in representatives:
            sales_by_representative = sales.filter(representative__id=representative['representative__id'])
            representative_group = {
                'representative':{
                    'code':representative['representative__code'],
                    'name':representative['representative__name']
                },
                'total_import':0,
                'total_import_mn':0,
                'dates':[]
            }
            for payment_method in payment_methods:
                representative_group["total_"+str(payment_method['id'])] = 0
            dates = sales_by_representative.saleGroupByDay()
            for date in dates:
                sales_by_date = sales_by_representative.filter(sale_date__date=date).order_by("sale_date")
                date_sale_group = {
                    'date':date,
                    'total_import':0,
                    'total_import_mn':0,
                    'sales':[]
                }
                for payment_method in payment_methods:
                    date_sale_group["total_"+str(payment_method['id'])] = 0
                for sale in sales_by_date:
                    if sale.status != "C":
                        sale_data = sale_subtotal(sale,self.with_out_tax)
                        sale_group = {
                            'sale':sale,
                            'totals':sale_data
                        }
                        date_sale_group['total_import'] += sale_data['total_num']
                        date_sale_group['total_import_mn'] += sale_data['total_mn_num']
                        for payment_method in payment_methods:
                            sale_group['payment_'+str(payment_method['id'])] = 0
                            amount = sale.sale_payments.amountAggregateByPaymentMethod(payment_method['id'])['amount_total']
                            sale_group['payment_'+str(payment_method['id'])] += amount * -1 if sale.status == "R" else amount
                            date_sale_group["total_"+str(payment_method['id'])] += amount * -1 if sale.status == "R" else amount
                        date_sale_group['sales'].append(sale_group)
                    else:
                        sale_group = {
                            'sale':sale,
                        }
                        date_sale_group['sales'].append(sale_group)
                for payment_method in payment_methods:
                    representative_group["total_"+str(payment_method['id'])] += date_sale_group["total_"+str(payment_method['id'])]
                representative_group['total_import'] += date_sale_group['total_import']
                representative_group['total_import_mn'] += date_sale_group['total_import_mn']
                representative_group['dates'].append(date_sale_group)
            for payment_method in payment_methods:
                context["total"]["total_"+str(payment_method['id'])] += representative_group["total_"+str(payment_method['id'])]
            context["total"]['total_import'] += representative_group['total_import']
            context["total"]['total_import_mn'] += representative_group['total_import_mn']
            context['representatives'].append(representative_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('sale_report_by_payment_method_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "Relacion de Cupones por Forma de Pago: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            total_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:Q', 20)
            worksheet.merge_range("A1:C1", self.property.name, merge_format)
            worksheet.merge_range("A2:C2", "Relacion de Cupones por Forma de Pago", merge_format)
            worksheet.merge_range("A3:B3", "Periodo: {} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.write("C3", "Fecha Emision: {}".format(date_time_now.strftime("%d/%m/%Y")), merge_format)
            column = 0
            if self.summary is not True:
                worksheet.write("B4", "Cupon", header_format)
                worksheet.write("C4", "Cve", header_format)
                worksheet.write("D4", "Servicio", header_format)
                worksheet.write("E4", "Grupo", header_format)
                worksheet.write("F4", "Habit", header_format)
                worksheet.write("G4", "Ingreso Neto", header_format)
                worksheet.write("H4", "IVA", header_format)
                worksheet.write("I4", "Ing.Dlls", header_format)
                worksheet.write("J4", "Ing.PESOS", header_format)
                worksheet.write("K4", "Desc", header_format)
                column = 11
            else:
                worksheet.write("B4", "Representante", header_format)
                worksheet.write("C4", "Ing.Dlls", header_format)
                worksheet.write("D4", "Ing.PESOS", header_format)
                column = 4
            for payment_method in payment_methods:
                worksheet.write(3,column, payment_method['name'], header_format)
                column+= 1
            row = 4
            for representative_group in context['representatives']:
                worksheet.merge_range(row,0,row,1, "{} - {} ".format(representative_group['representative']['code'],representative_group['representative']['name']),header_format)
                row += 1
                for date_sale_group in representative_group['dates']:
                    if self.summary is not True:
                        worksheet.write(row,0, date_sale_group['date'].strftime("%d/%m/%Y"))
                        row += 1
                        for sale_data in date_sale_group['sales']:
                            sale = sale_data['sale']
                            if sale.status != "C":
                                totals = sale_data['totals']
                                worksheet.write(row,1, "R{}".format(str(sale.sale_key).zfill(8)) if sale.status == "R" else str(sale.sale_key).zfill(8))
                                worksheet.write(row,2, str(sale.service.id).zfill(4))
                                worksheet.write(row,3, sale.service.name)
                                if sale.service.business_group is not None:
                                    worksheet.write(row,4, sale.service.business_group.name)
                                worksheet.write(row,5, sale.room)
                                worksheet.write(row,6, totals['total_no_tax'])
                                worksheet.write(row,7, totals['total_only_tax'])
                                worksheet.write(row,8, totals['total_num'])
                                worksheet.write(row,9, totals['total_mn_num'])
                                worksheet.write(row,10, totals['discount_num'])
                                column_ = 11
                                for payment_method in payment_methods:
                                    worksheet.write(row,column_, sale_data["payment_"+str(payment_method['id'])])
                                    column_+= 1
                            else:
                                worksheet.write(row,1, "C{}".format(str(sale.sale_key).zfill(8)))
                                worksheet.write(row,5, sale.room)
                                worksheet.merge_range(row,11,row,column, "*** CANCELADO",merge_format)
                            row += 1
                    worksheet.write(row,0, date_sale_group['date'].strftime("%d/%m/%Y"))
                    if self.summary is not True:
                        column_ = 8
                    else:
                        column_ = 2
                    worksheet.write(row,column_, date_sale_group['total_import'], total_format)
                    column_+= 1
                    worksheet.write(row,column_, date_sale_group['total_import_mn'], total_format)
                    column_+= 2
                    for payment_method in payment_methods:
                        worksheet.write(row,column_, date_sale_group["total_"+str(payment_method['id'])], total_format)
                        column_+= 1
                    row += 1
                worksheet.merge_range(row,0,row,1, "{} - {} ".format(representative_group['representative']['code'],representative_group['representative']['name']),total_format)
                if self.summary is not True:
                    column_ = 8
                else:
                    column_ = 2
                worksheet.write(row,column_, representative_group['total_import'], total_format)
                column_+= 1
                worksheet.write(row,column_, representative_group['total_import_mn'], total_format)
                column_+= 2
                for payment_method in payment_methods:
                    worksheet.write(row,column_, representative_group["total_"+str(payment_method['id'])], total_format)
                    column_+= 1
                row += 1
            if self.summary is not True:
                column_ = 8
            else:
                column_ = 2
            worksheet.write(row,column_, context['total']['total_import'], total_format)
            column_+= 1
            worksheet.write(row,column_, context['total']['total_import_mn'], total_format)
            column_+= 2
            for payment_method in payment_methods:
                worksheet.write(row,column_, context['total']["total_"+str(payment_method['id'])], total_format)
                column_+= 1
            row += 1
            workbook.close()
            output.seek(0)
            filename = 'Relacion de Cupones por Forma de Pago {} .xlsx'.format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_sale_by_sale_type_and_hotel(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'sale_types': [],
            'total_pax':sales.filter(status__in=["A","R"]).paxTotal(),
            'total_import':0,
        }
        sale_type_ids = sales.order_by('sale_type').values_list('sale_type', flat=True).distinct('sale_type')
        for sale_type_id in sale_type_ids:
            print()
            sale_type = SaleType.objects.get(id=sale_type_id)
            sales_by_sale_types = sales.filter(status__in=["A","R"],sale_type=sale_type)
            sale_type_group = {
                'name':sale_type.name,
                'total_pax':sales_by_sale_types.paxTotal(),
                'total_import':0,
                'hotels': [],
            }
            hotel_ids = sales_by_sale_types.order_by('hotel').values_list('hotel', flat=True).distinct('hotel')
            for hotel_id in hotel_ids:
                hotel = Hotel.objects.get(id=hotel_id)
                sales_by_hotel = sales_by_sale_types.filter(status__in=["A","R"],hotel=hotel)
                hotel_group = {
                    'name':hotel.name,
                    'total_pax':sales_by_hotel.paxTotal(),
                    'total_import':0,
                }
                for sale in sales_by_hotel:
                    totals = sale_subtotal(sale)
                    hotel_group['total_import'] += totals['total_num']
                sale_type_group['hotels'].append(hotel_group)
                sale_type_group['total_import'] += hotel_group['total_import']
            context['total_import'] += sale_type_group['total_import']
            context['sale_types'].append(sale_type_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_sale_by_sale_type_and_hotel_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "DETALLES DE VENTAS Y COMISIONES: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            provider_format = workbook.add_format({
                'bold': True,
                "align": "right",
                "valign": "vcenter",
                'font_size':12,
            })
            total_format = workbook.add_format({
                'top':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:E2", "DETALLES DE VENTAS Y COMISIONES", merge_format)
            worksheet.merge_range("C3:E3", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("F1:G2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            row = 3
            for sale_type_group in context['sale_types']:
                worksheet.merge_range(row,3,row,4, sale_type_group['name'], merge_format)
                row += 1
                worksheet.merge_range(row,0,row,4, "HOTEL", header_format)
                worksheet.write(row,5, "PAX", header_format)
                worksheet.write(row,6, "IMPORTE", header_format)
                row += 1
                for hotel in sale_type_group['hotels']:
                    worksheet.merge_range(row,0,row,4, hotel['name'])
                    worksheet.write(row,5, "{}.{}".format(int(hotel['total_pax']['adults_total']),int(hotel['total_pax']['childs_total'])))
                    worksheet.write(row,6, hotel['total_import'])
                    row += 1
                worksheet.write(row,5, "{}.{}".format(int(hotel['total_pax']['adults_total']),int(hotel['total_pax']['childs_total'])), total_format)
                worksheet.write(row,6, hotel['total_import'], total_format)
                row += 1
            worksheet.merge_range(row,0,row,4, "GRAN TOTAL", provider_format)
            worksheet.write(row,5, "{}.{}".format(int(context['total_pax']['adults_total']),int(context['total_pax']['childs_total'])), total_format)
            worksheet.write(row,6, context['total_import'], total_format)
            workbook.close()
            output.seek(0)
            filename = "DETALLES DE VENTAS Y COMISIONES: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def report_refund_by_representatives(self,sales):
        sales = sales.filter(status='R')
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'representatives': [],
            'total_import':0,
        }
        sales = sales.order_by('sale_key')
        representative_ids = sales.order_by('representative').values_list('representative', flat=True).distinct('representative')
        for representative_id in representative_ids:
            representative = Representative.objects.get(id=representative_id)
            sales_by_representative = sales.filter(representative=representative)
            representative_group = {
                'name':representative.name,
                'total_import':0,
                'sales': []
            }
            for sale in sales_by_representative:
                totals = sale_subtotal(sale)
                sale_group = {
                    'sale':sale,
                    'import':totals['total_num']
                }
                representative_group['total_import'] += totals['total_num']
                representative_group['sales'].append(sale_group)
            context['total_import'] += representative_group['total_import']
            context['representatives'].append(representative_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_refund_by_representatives_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "REPORTE DE CUPONES REEMBOLSADOS: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            representative_format = workbook.add_format({
                'bold': True,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:E2", "REPORTE DE CUPONES REEMBOLSADOS", merge_format)
            worksheet.merge_range("F1:G2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)

            worksheet.write("A3", "FECHA", header_format)
            worksheet.write("B3", "CUPON", header_format)
            worksheet.merge_range("C3:E3", "SERVICIO", header_format)
            worksheet.write("F3", "IMPORTE", header_format)
            row = 3
            for representative_group in context['representatives']:
                worksheet.merge_range(row,0,row,1, representative_group['name'],representative_format)
                row += 1
                for sale_group in representative_group['sales']:
                    sale = sale_group['sale']
                    worksheet.write(row,0, sale.sale_date.strftime("%d/%m/%Y"))
                    worksheet.write(row,1,str(sale.sale_key).zfill(8))
                    worksheet.merge_range(row,2,row,4,sale.service.name)
                    worksheet.write(row,5, sale_group['import'])
                    row += 1
                worksheet.merge_range(row,0,row,5, representative_group['total_import'],totals_format)
                worksheet.write(row,6, "{} Cupones".format(len(representative_group['sales'])))
                row += 1
            worksheet.merge_range(row,0,row,5, context['total_import'],totals_format)
            worksheet.write(row,6, "{} Cupones".format(len(context['representatives'])))
            workbook.close()
            output.seek(0)
            filename = "REPORTE DE CUPONES REEMBOLSADOS: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
    
    def report_sales_with_discount(self,sales):
        sales = sales.filter(discount__gt=0)
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'total_sales':sales.count(),
            'representatives': [],
            'total_num':0,
            'discount_num':0,
            'subtotal_num':0,
        }
        sales = sales.order_by('sale_key')
        representative_ids = sales.order_by('representative').values_list('representative', flat=True).distinct('representative')
        for representative_id in representative_ids:
            representative = Representative.objects.get(id=representative_id)
            sales_by_representative = sales.filter(representative=representative)
            representative_group = {
                'name':representative.name,
                'total_import':0,
                'sales': [],
                'total_num':0,
                'discount_num':0,
                'subtotal_num':0,
            }
            for sale in sales_by_representative:
                sale_group = {
                    'sale':sale,
                }
                sale_group.update(sale_subtotal(sale,self.with_out_tax))
                representative_group['total_num'] += sale_group['total_num']
                representative_group['discount_num'] += sale_group['discount_num']
                representative_group['subtotal_num'] += sale_group['subtotal_num']
                representative_group['sales'].append(sale_group)
            context['total_num'] += representative_group['total_num']
            context['discount_num'] += representative_group['discount_num']
            context['subtotal_num'] += representative_group['subtotal_num']
            context['representatives'].append(representative_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_sales_with_discount_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "REPORTE DE CUPONES CON DESCUENTO: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            representative_format = workbook.add_format({
                'bold': True,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:E2", "REPORTE DE CUPONES CON DESCUENTO", merge_format)
            worksheet.merge_range("F1:G2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)

            worksheet.write("A3", "FECHA", header_format)
            worksheet.write("B3", "CUPON", header_format)
            worksheet.merge_range("C3:E3", "SERVICIO", header_format)
            worksheet.write("F3", "Vta Bruta", header_format)
            worksheet.write("G3", "Descuento", header_format)
            worksheet.write("H3", "%", header_format)
            worksheet.write("I3", "Vta Neta", header_format)
            row = 3
            for representative_group in context['representatives']:
                worksheet.merge_range(row,0,row,1, representative_group['name'],representative_format)
                row += 1
                for sale_group in representative_group['sales']:
                    sale = sale_group['sale']
                    worksheet.write(row,0, sale.sale_date.strftime("%d/%m/%Y"))
                    worksheet.write(row,1,str(sale.sale_key).zfill(8))
                    worksheet.merge_range(row,2,row,4,sale.service.name)
                    worksheet.write(row,5, sale_group['subtotal_num'])
                    worksheet.write(row,6, sale_group['discount_num'])
                    worksheet.write(row,7, sale.discount)
                    worksheet.write(row,8, sale_group['total_num'])
                    row += 1
                worksheet.write(row,0,"",totals_format)
                worksheet.write(row,1,"{} Cupones".format(len(representative_group['sales'])),totals_format)
                worksheet.merge_range(row,2,row,4,"",totals_format)
                worksheet.write(row,5, representative_group['subtotal_num'],totals_format)
                worksheet.write(row,6, representative_group['discount_num'],totals_format)
                worksheet.write(row,7, "",totals_format)
                worksheet.write(row,8, representative_group['total_num'],totals_format)
                row += 1
            worksheet.write(row,0,"",totals_format)
            worksheet.write(row,1,"{} Cupones".format(context['total_sales']),totals_format)
            worksheet.merge_range(row,2,row,4,"",totals_format)
            worksheet.write(row,5, context['subtotal_num'],totals_format)
            worksheet.write(row,6, context['discount_num'],totals_format)
            worksheet.write(row,7, "",totals_format)
            worksheet.write(row,8, context['total_num'],totals_format)
            workbook.close()
            output.seek(0)
            filename = "REPORTE DE CUPONES CON DESCUENTO: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response

    def report_sales_by_representatives_and_providers(self,sales):
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'total_pax':        sales.filter(status__in=["A","R"]).paxTotal(),
            'total_import':0,
            'total_comission':0,
            'providers': [],
        }
        provider_ids = sales.order_by('service__provider').values_list('service__provider', flat=True).distinct('service__provider')
        for provider_id in provider_ids:
            provider = Provider.objects.get(id=provider_id)
            sales_by_provider = sales.filter(status__in=["A","R"],service__provider=provider)
            provider_group = {
                'name':provider.name,
                'total_pax':sales_by_provider.paxTotal(),
                'total_import':0,
                'total_comission':0,
                'services': []
            }
            service_ids = sales_by_provider.order_by('service').values_list('service', flat=True).distinct('service')
            for service_id in service_ids:
                service = Service.objects.get(id=service_id)
                sales_by_service = sales_by_provider.filter(service=service)
                service_group = {
                    'name':service.name,
                    'total_pax':sales_by_service.paxTotal(),
                    'total_import':0,
                    'total_cost':0,
                    'total_comission':0,
                }
                for sale in sales_by_service:
                    totals = sale_subtotal(sale,self.with_out_tax)
                    service_group['total_import'] += totals['total_num']
                    service_group['total_comission'] += totals['rep_comission']
                provider_group['services'].append(service_group)
                provider_group['total_import'] += service_group['total_import']
                provider_group['total_comission'] += service_group['total_comission']
            context['providers'].append(provider_group)
            context['total_import'] += provider_group['total_import']
            context['total_comission'] += provider_group['total_comission']
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_sales_by_representatives_and_providers_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "VENTAS POR REPRESENTANTE Y PROVEEDOR: {} .pdf".format(self.date_filter,date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            provider_format = workbook.add_format({
                'bold': True,
                "align": "right",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 15)
            worksheet.merge_range("A1:B1", self.property.name, merge_format)
            worksheet.merge_range("C1:E1", "VENTAS POR REPRESENTANTE Y PROVEEDOR", merge_format)
            worksheet.merge_range("C2:E2", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("F1:G1", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            worksheet.merge_range("A3:C3", "SERVICIO", header_format)
            worksheet.merge_range("D3:E3", "PROVEEDOR", header_format)
            worksheet.write("F4", "PAX", header_format)
            worksheet.write("G4", "IMPORTE", header_format)
            worksheet.write("H4", "COMISION", header_format)
            row = 4
            for provider in context['providers']:
                worksheet.merge_range(row,0,row,1,provider['name'],header_format)
                row += 1
                for service in provider['services']:
                    worksheet.merge_range(row,0,row,2, service['name'])
                    worksheet.merge_range(row,3,row,4, provider['name'])
                    worksheet.write(row,5, "{}.{}".format(int(service['total_pax']['adults_total']),int(service['total_pax']['childs_total'])))
                    worksheet.write(row,6, service['total_import'])
                    worksheet.write(row,7, service['total_comission'])
                    row += 1
                worksheet.write(row,5, "{}.{}".format(int(provider['total_pax']['adults_total']),int(provider['total_pax']['childs_total'])),totals_format)
                worksheet.write(row,6, provider['total_import'],totals_format)
                worksheet.write(row,7, provider['total_comission'],totals_format)
                row += 1
            workbook.close()
            output.seek(0)
            filename = "VENTAS POR REPRESENTANTE Y PROVEEDOR: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
    
    def report_sales_pax_by_services(self,sales):
        from SalesApp.models import Sale
        date_time_now = datetime.now()
        context = {
            'date':             date_time_now,
            'user':             self.request.user,
            'start_date':       self.start_date,
            'due_date':         self.due_date,
            'property':         self.property,
            'services': [],
        }
        
        sales = sales.saleRefundMark().filter(has_refund=False).exclude(status="C")
        service_ids = sales.order_by('service').values_list('service', flat=True).distinct('service')
        for service_id in service_ids:
            service = Service.objects.get(id=service_id)
            sales_by_service = sales.filter(service=service)
            service_group = {
                'name':service.name,
                'provider':service.provider.name,
                'dates':[]
            }
            dates = sales_by_service.serviceGroupByDay()
            for date in dates:
                sales_by_date = sales_by_service.filter(service_date=date).order_by('sale_key')
                date_sale_group = {
                    'date':date,
                    'pax_total':sales_by_date.paxTotal(),
                    'sales':[]
                }
                for sale in sales_by_date:
                    date_sale_group['sales'].append(sale)
                service_group['dates'].append(date_sale_group)
            context['services'].append(service_group)
        print(len(connection.queries))
        if self.format == "pdf":
            template = get_template('report_sales_pax_by_services_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "LISTA DE PAX PARA EXCURSIONES: {} .pdf".format(self.date_filter,date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'top':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            subheader_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "right",
                "valign": "vcenter",
                'font_size':12,
            })
            subheader_format_text = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:H', 12)
            worksheet.merge_range("A1:B1", self.property.name, merge_format)
            worksheet.merge_range("C1:E1", "LISTA DE PAX PARA EXCURSIONES", merge_format)
            worksheet.merge_range("C2:E2", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
            worksheet.merge_range("F1:G1", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            row = 3
            for service_group in context['services']:
                for date_sale_group in service_group['dates']:
                    worksheet.merge_range(row,0,row,1, date_sale_group['date'].strftime("%d/%m/%Y"), header_format)
                    worksheet.merge_range(row,2,row,4, service_group['name'], header_format)
                    worksheet.write(row,5, "Prov.", header_format)
                    worksheet.merge_range(row,6,row,8, service_group['provider'], header_format)
                    row += 1
                    worksheet.write(row,1, "CUPON",subheader_format)
                    worksheet.merge_range(row,2,row,4, "HOTEL",subheader_format_text)
                    worksheet.write(row,5, "HAB",subheader_format)
                    worksheet.write(row,6, "PAX",subheader_format)
                    worksheet.merge_range(row,7,row,9, "NOMBRE PAX",subheader_format_text)
                    row += 1
                    for sale in date_sale_group['sales']:
                        worksheet.write(row,1, str(sale.sale_key).zfill(8))
                        worksheet.merge_range(row,2,row,4, sale.hotel.name)
                        worksheet.write(row,5, sale.room)
                        worksheet.write(row,6, "{}.{}".format(int(sale.adults),int(sale.childs)))
                        worksheet.merge_range(row,7,row,9, sale.name_pax)
                        row += 1
                    worksheet.write(row,6, "{}.{}".format(int(date_sale_group['pax_total']['adults_total']),int(date_sale_group['pax_total']['childs_total'])),totals_format)
            workbook.close()
            output.seek(0)
            filename = "LISTA DE PAX PARA EXCURSIONES: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response

class SaleSAPReportController():
    def __init__(self, request, start_date, due_date, property, transfer_service, sap_codes, currency, date_filter, with_out_tax):
        self.request = request
        self.start_date = start_date
        self.due_date = due_date
        self.property = property
        self.transfer_service = transfer_service
        self.sap_codes = sap_codes
        self.currency = currency
        self.date_filter = date_filter
        self.with_out_tax = with_out_tax

    def filters(self):
        from django.db.models import Q
        sales = models.Sale.objects.tableOptimization()
        if self.date_filter == "xFechaVenta":
            sales = sales.saleFilterSaleDateRange(
                self.start_date,self.due_date,self.property
            ).order_by("sale_date")
        elif self.date_filter == "xFechaServicio":
            sales = sales.saleFilterServiceDateRange(
                self.start_date,self.due_date,self.property
            ).order_by("service_date")
        else:
            raise CustomValidation("No hay filtro para busqueda de fechas", 'date_filter', status.HTTP_400_BAD_REQUEST)
        if self.transfer_service is True:
            sales = sales.filter(Q(service__is_transfer=True)|Q(service__business_group__name="Transportación"))
        if isinstance(self.sap_codes, list):
            sales = sales.saleFilterBySapCode(self.sap_codes)
        elif self.sap_codes == "RCD":
            sales = sales.saleFilterBySapCode(["RCD",])
        elif self.sap_codes == "-RCD":
            sales = sales.saleExcludeBySapCode(["RCD",])
        if self.currency == "MN" or self.currency == "USD":
            sales = sales.filter(service__provider__currency=self.currency)
        return sales
    
    def get_context(self):
        sales = self.filters()
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output)
        worksheet_sales = workbook.add_worksheet('Ventas')
        worksheet_sales = self.report_sap_sheet(worksheet_sales, sales.filter(status="A"))
        worksheet_refunds = workbook.add_worksheet('Reembolsos')
        worksheet_refunds = self.report_sap_sheet(worksheet_refunds, sales.filter(status="R"))
        workbook.close()
        output.seek(0)
        filename = "{}-{}-report-sap-sales-{}-{}.xlsx".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        response = HttpResponse(
            output.getvalue(),
            content_type='application/ms-excel;charset=utf-8'
        )
        response['Content-Disposition'] = "attachment; filename={}".format(filename)
        response["Content-Encoding"] = "UTF-8"
        return response
    
    def report_sap_sheet(self, worksheet, sales):
        C = 0
        P = 1
        S = 2
        for sale in sales:
            sale_data = sale_subtotal(sale,self.with_out_tax)
            if sale_data['total_cost_num'] != 0:
                worksheet.write(C, 0, "C")

                worksheet.write(P, 0, "P")

                worksheet.write(S, 0, "S")

                worksheet.write(C, 1, "ZOPE")
                worksheet.write(C, 3, sale.service.provider.sap_code)
                worksheet.write(C, 4, sale.sale_date.date().strftime("%d.%m.%Y"))
                if "Punta Cana" in self.property.name:
                    worksheet.write(C, 5, "OCDO")
                else: 
                    worksheet.write(C, 5, "OCMX")
                worksheet.write(C, 6, "OPE")
                worksheet.write(C, 7, str(sale.sale_type.department_cecos.code)[0:4] if sale.sale_type.department_cecos is not None else "")

                worksheet.write(P, 9, "K")
                worksheet.write(P, 10, "F")
                worksheet.write(P, 11, sale.service.name)
                worksheet.write(P, 14, sale.service_date.strftime("%d.%m.%Y"))
                worksheet.write(P, 16, "SERV015")
                worksheet.write(P, 17, str(sale.sale_type.department_cecos.code)[0:4] if sale.sale_type.department_cecos is not None else "")
                worksheet.write(P, 18, sale.sale_type.warehouse_code if sale.sale_type.warehouse_code is not None else "")

                worksheet.write(S, 19, sale.sale_type.sap_code)
                worksheet.write(S, 21, 1)

                
                if sale.service.provider.currency == "MN":
                    exchange_rate = ExchangeRate.objects.filter(start_date__lte=sale.sale_date.date(),provider=sale.service.provider).order_by("-start_date").first()
                    if exchange_rate is not None:
                        sale_cost = sale_data['total_cost_num']*exchange_rate.usd_currency
                        worksheet.write(S, 23, sale_cost)
                    else:
                        exchange_rate = ExchangeRate.objects.filter(start_date__lte=sale.sale_date.date(),type='SALE',property=sale.property).order_by("-start_date").first()
                        sale_cost = sale_data['total_cost_num']*exchange_rate.usd_currency
                        worksheet.write(S, 23, sale_cost)
                    worksheet.write(S, 31, "MXN")
                else:
                    worksheet.write(S, 23, sale_data['total_cost_num'])
                    worksheet.write(S, 31, "USD")
                worksheet.write(S, 25, sale.sale_type.department_cecos.code if sale.sale_type.department_cecos is not None else "")

                worksheet.write(C, 26, sale.service.name)
                worksheet.write(C, 27, sale.sale_date.date().strftime("%d.%m.%Y"))
                worksheet.write(C, 28, sale.service_date.strftime("%d.%m.%Y"))
                worksheet.write(C, 29, str(sale.sale_key).zfill(8))
                C+=3
                P+=3
                S+=3
        worksheet.autofit()
        return worksheet
        
class OperatorReportController():
    def __init__(self, request, date, availability_group, property):
        self.request = request
        self.date = date
        self.property = property
        self.availability_group = availability_group

    def get_context(self):
        from django.db.models import Q
        from SalesApp.views import schedule_by_datetime
        template = get_template('operation_vp_report.html')
        context = {
            'host'              :   self.request.get_host(),
            'environment'       :   getattr(serverconfig,"environment","http"),
            'date'              :   self.date,
            'today'             :   datetime.now(),
            'user'              :   self.request.user,
            'availability_group':   self.availability_group,
            'property'          :   self.property,
            'unit_groups'       :   []
        }
        group = Group.objects.filter(availability_groups=self.availability_group).first()
        if group is not None and self.availability_group is not None:
            sales = models.Sale.objects.tableOptimization().saleFilterServiceDate(self.date,self.property).filter(service__availability_group__group=group)
        elif self.availability_group is not None:    
            sales = models.Sale.objects.tableOptimization().saleFilterServiceDate(self.date,self.property).filter(service__availability_group=self.availability_group)
        else:    
            sales = models.Sale.objects.tableOptimization().saleFilterServiceDate(self.date,self.property)
        sales_by_no_unit = sales.filter(Q(unit='')|Q(unit__isnull=True),service__availability_group__isnull=False)
        if sales_by_no_unit.exists():
            unit_group = {
                'unit':"Sin unidad",
                'availability_groups':[]
            }
            availability_groups = sales_by_no_unit.order_by('service__availability_group').values_list('service__availability_group', flat=True).distinct('service__availability_group')
            for availability_group in availability_groups:
                availability_group_ = AvailabilityGroup.objects.get(id=availability_group)
                sales_by_availability_group = sales_by_no_unit.filter(service__availability_group__id=availability_group)
                times = sales_by_availability_group.exclude(time__isnull=True).order_by('time').values_list('time', flat=True).distinct('time')
                for time in times:
                    print(time)
                    schedule = schedule_by_datetime(self.date,availability_group_,time)
                    if schedule is not None:
                        sales_by_time = sales_by_availability_group.filter(time=time).saleRefundMark().filter(has_refund=False).order_by('service__name')
                        if sales_by_time.exists():
                            unit_availability_group = {
                                'availability_group':availability_group_,
                                'time':time,
                                'sales':serializers.SaleWithPickUpSerializer(sales_by_time,many=True).data
                            }
                            unit_availability_group.update(sales_by_time.paxTotal())
                    unit_group['availability_groups'].append(unit_availability_group)
            context['unit_groups'].append(unit_group)
        units = sales.exclude(Q(unit='')|Q(unit__isnull=True)).order_by('unit').values_list('unit', flat=True).distinct('unit')
        for unit in units:
            unit_group = {
                'unit':unit,
                'availability_groups':[]
            }
            sales_by_unit = sales.filter(unit=unit)
            availability_groups = sales_by_unit.order_by('service__availability_group').values_list('service__availability_group', flat=True).distinct('service__availability_group')
            for availability_group in availability_groups:
                availability_group_ = AvailabilityGroup.objects.get(id=availability_group)
                sales_by_availability_group = sales_by_unit.filter(service__availability_group__id=availability_group)
                times = sales_by_availability_group.exclude(time__isnull=True).order_by('time').values_list('time', flat=True).distinct('time')
                for time in times:
                    schedule = schedule_by_datetime(self.date,availability_group_,time)
                    if schedule is not None:
                        sales_by_time = sales_by_availability_group.filter(time=time).saleRefundMark().filter(has_refund=False).order_by('service__name')
                        if sales_by_time.exists():
                            unit_availability_group = {
                                'availability_group':availability_group_,
                                'time':time,
                                'sales':serializers.SaleWithPickUpSerializer(sales_by_time,many=True).data
                            }
                            unit_availability_group.update(sales_by_time.paxTotal())
                    unit_group['availability_groups'].append(unit_availability_group)
            context['unit_groups'].append(unit_group)
        html  = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            response = HttpResponse(result.getvalue(), content_type='application/pdf')
            filename =  "{}-operation-report-vp-{}-{}.pdf".format(str(self.date),self.property.code,str(datetime.now()))
            content = "inline; filename={}".format(filename)
            response['Content-Disposition'] = content
            return response
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
    
class ServiceListController():
    def __init__(self, request, date, print_due, property, type, format, services, providers):
        from GeneralApp.models import ExchangeRate
        self.request = request
        self.date = date
        self.print_due = print_due
        self.property = property
        self.type = type
        self.format = format
        self.services = services
        self.providers = providers

    def filters(self):
        service_rates = models.ServiceRate.objects.dateSimple(self.date)
        if len(self.services) > 0:
            service_rates = service_rates.filter(service__id__in=self.services)
        else:
            services = Service.objects.filter(properties=self.property).distinct()
            service_rates = service_rates.filter(service__in=services)
        if len(self.providers) > 0:
            service_rates = service_rates.filter(service__provider__id__in=self.providers)
        return service_rates
    
    def get_context(self):
        service_rates = self.filters()
        providers = service_rates.order_by('service__provider').values_list('service__provider', flat=True).distinct('service__provider')
        payment_types_ids = service_rates.order_by('service_rate_comissions__payment_type').values_list('service_rate_comissions__payment_type', flat=True).distinct('service_rate_comissions__payment_type')
        payment_types = models.PaymentType.objects.filter(id__in=payment_types_ids)
        date_time_now = datetime.now()
        context = {
            'date_print':             date_time_now,
            'user':             self.request.user,
            'date':             self.date,
            'type':             self.type,
            'print_due':        self.print_due,
            'property':         self.property,
            'payment_types':    payment_types,
            'service_rates':    []
        }
        for provider in providers:
            service_rates_by_provider = service_rates.filter(service__provider__id=provider)
            provider_group = {
                'provider':ProviderSerializer(Provider.objects.get(id=provider)).data,
                'service_rates':[]
            }
            for service_rate in service_rates_by_provider:
                service_rate_data = {
                    'cve':service_rate.service.id,
                    'name':service_rate.service.name,
                    'start_date':service_rate.start_date,
                    'due_date':service_rate.due_date,
                    'currency':service_rate.currency,
                }
                service_rate_data['adult_price'] = service_rate.adult_price
                service_rate_data['child_price'] = service_rate.child_price
                service_rate_data['adult_fee'] = service_rate.adult_price-(service_rate.adult_price*(service_rate.hard_rock_comission_adult/100))
                service_rate_data['child_fee'] = service_rate.child_price-(service_rate.child_price*(service_rate.hard_rock_comission_child/100))
                service_rate_data['adult_hc'] =  int(service_rate.hard_rock_comission_adult)
                service_rate_data['child_hc'] =  int(service_rate.hard_rock_comission_child)
                for payment_type in payment_types:
                    service_rate_comission = service_rate.service_rate_comissions.filter(payment_type=payment_type).first()
                    if service_rate_comission is not None:
                        service_rate_data["payment_type_"+str(payment_type.id)] = int(service_rate_comission.comission)
                    else:
                        service_rate_data["payment_type_"+str(payment_type.id)] = None
                provider_group['service_rates'].append(service_rate_data)
            context['service_rates'].append(provider_group)
        if self.format == "pdf":
            template = get_template('service_report_template.html')
            html  = template.render(context)
            result = BytesIO()
            pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
            if not pdf.err:
                response = HttpResponse(result.getvalue(), content_type='application/pdf')
                filename = "CATALOGO DE SERVICIOS POR PROVEEDOR: {} .pdf".format(date_time_now)
                content = "inline; filename={}".format(filename)
                response['Content-Disposition'] = content
                return response
            raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        elif self.format == "excel":
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            merge_format = workbook.add_format({
                "bold": 1,
                "align": "center",
                "valign": "vcenter",
            })
            header_format = workbook.add_format({
                'bold': True,
                'bottom':1,
                "align": "center",
                "valign": "vcenter",
                'font_size':12,
            })
            totals_format = workbook.add_format({
                'top':1,
                "valign": "vcenter",
                'font_size':12,
            })
            worksheet = workbook.add_worksheet()
            worksheet.set_column('A:N', 15)
            worksheet.merge_range("A1:B2", self.property.name, merge_format)
            worksheet.merge_range("C1:E2", "CATALOGO DE SERVICIOS POR PROVEEDOR", merge_format)
            worksheet.merge_range("F1:G2", "{}\n{}".format(self.request.user.username,date_time_now.strftime("%d/%m/%Y %H:%M:%S")), merge_format)
            worksheet.merge_range("A3:D3", "Esta edición cancela todas la anteriores {}".format(context['date'].strftime("%B %d, %Y")), merge_format)
            worksheet.write("A4", "CVE", header_format)
            worksheet.merge_range("B4:C4", "SERVICIO", header_format)
            column = 3
            if self.print_due is True:
                worksheet.write(3,column, "Fecha Inicio", header_format)
                column += 1
                worksheet.write(3,column, "Fecha Fin", header_format)
                column += 1
            worksheet.write(3,column, "$ADULT", header_format)
            column += 1
            worksheet.write(3,column, "$MENOR", header_format) 
            column += 1
            worksheet.write(3,column, "Mon", header_format)
            column += 1
            if self.type != "exclude_commission" and self.type != "just_commission":
                worksheet.write(3,column, "COM A G ADULTO", header_format)
                column += 1
                worksheet.write(3,column, "COM A G MENOR", header_format)
                column += 1
            if self.type != "exclude_commission":
                for payment_type in context['payment_types']:
                    worksheet.write(3,column, payment_type.name, header_format)
                    column += 1
            if self.type != "just_commission":
                worksheet.write(3,column, "Costo Adulto", header_format)
                column += 1
                worksheet.write(3,column, "Costo Menor", header_format)
                column += 1
            row = 4
            for service_rate_provider in context['service_rates']:
                worksheet.merge_range(row,0,row,2, "PROVEEDOR: {}".format(service_rate_provider['provider']['name']))
                worksheet.merge_range(row,3,row,4, "TELEFONO: {}".format(service_rate_provider['provider']['phone']))
                row += 1
                for service_rate in service_rate_provider['service_rates']:
                    worksheet.write(row,0, service_rate['cve'])
                    worksheet.merge_range(row,1,row,2, service_rate['name'])
                    column = 3
                    if self.print_due is True:
                        worksheet.write(row,column, service_rate['start_date'].strftime("%d/%m/%Y"))
                        column += 1
                        worksheet.write(row,column, service_rate['due_date'].strftime("%d/%m/%Y"))
                        column += 1
                    worksheet.write(row,column, service_rate['adult_price'])
                    column += 1
                    worksheet.write(row,column, service_rate['child_price'])
                    column += 1
                    worksheet.write(row,column, service_rate['currency'])
                    column += 1
                    if self.type != "exclude_commission" and self.type != "just_commission":
                        worksheet.write(row,column, service_rate['adult_hc'])
                        column += 1
                        worksheet.write(row,column, service_rate['adult_hc'])
                        column += 1
                    if self.type != "exclude_commission":
                        for payment_type in context['payment_types']:
                            value = service_rate.get('payment_type_'+str(payment_type.id),0)
                            worksheet.write(row,column, value if value is not None else 0)
                            column += 1
                    if self.type != "just_commission":
                        worksheet.write(row,column, service_rate['adult_fee'])
                        column += 1
                        worksheet.write(row,column, service_rate['child_fee'])
                        column += 1
                    row += 1
            workbook.close()
            output.seek(0)
            filename = "CATALOGO DE SERVICIOS POR PROVEEDOR: {} .xlsx".format(date_time_now)
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response


            