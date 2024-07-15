from GeneralApp.models                  import Service, Provider, Property, TRANSFER_TYPE_CHOICES, TRANSFER_TYPE_CHOICES_SIMPLE, Unit, Hotel, ExchangeRate
from OperationsApp                      import serializers, models, querysets, views
from GeneralApp.utils                   import CustomValidation
from rest_framework                     import viewsets, permissions, status
from django.template.loader             import get_template
from django.http                        import HttpResponse,JsonResponse
from rest_framework.response            import Response
from datetime                           import time, date, datetime, timezone, timedelta
from io                                 import BytesIO
from xhtml2pdf                          import pisa
from django.conf                        import settings
import xlsxwriter
from docx                               import Document
from docx.shared                        import Inches, Pt
from docx.enum.text                     import WD_ALIGN_PARAGRAPH
from docx.enum.section                  import WD_ORIENT
from operacionesVP import serverconfig

def daterange(start_date, due_date):
    due_date = due_date + timedelta(1)
    for n in range(int((due_date - start_date).days)):
        yield start_date + timedelta(n)

class OperatorReportController():
    def __init__(self, request, date, property, type, unit, hotels):
        self.request = request
        self.date = date
        self.property = property
        self.type = type
        self.unit = unit
        self.hotels = hotels
        self.hotels_name = "Todos"

    def filters(self):
        reservation_services = models.ReservationService.objects.tableOptimization(
        ).operationFilter(
            self.date,self.property
        ).exclude(
            reservation__status="CANCEL"
        ).commentsAnnotate(
        ).order_by_operation_time()
        if self.unit != "":
            reservation_services = reservation_services.operationFilterByUnit(self.unit)
        if len(self.hotels) > 0:
            reservation_services = reservation_services.operationFilterByHotels(self.hotels)
        return reservation_services
    
    def units_agroups_report(self,reservation_services):
        reservation_services = reservation_services.filter(unit_value__isnull=False).order_by_unit_operation_time()
        units = []
        for reservation_service in reservation_services:
            filtered = filter(lambda element: reservation_service.unit_value == element, units)
            if len(list(filtered)) == 0:
                units.append(reservation_service.unit_value)
        units.append(None)
        return units

    def get_context(self):
        reservation_services = self.filters()
        units_group = self.units_agroups_report(reservation_services)
        reservation_services_groups = []
        for unit_value in units_group:
            reservation_services_group_query = reservation_services.filter(unit_value=unit_value)
            if reservation_services_group_query.exists():
                reservation_services_group = {
                    'reservation_services':reservation_services_group_query,
                }
                reservation_services_group.update(reservation_services_group_query.paxAggregate())
                reservation_services_groups.append(reservation_services_group)
        if self.type == "pdf":
            result, filename = self.report_in_pdf(reservation_services_groups,reservation_services)
            response = HttpResponse(result.getvalue(), content_type='application/pdf')
            content = "inline; filename={}".format(filename)
            response['Content-Disposition'] = content
            return response
        else:
            result, filename = self.report_in_excel(reservation_services_groups)
            response = HttpResponse(
                result,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename=%s' % filename
            return response
        
    def get_context_for_providers(self, providers_ids=[]):
        reservation_services = self.filters()
        reservation_services = reservation_services.filter(unit__isnull=False)
        if len(providers_ids) > 0:
            providers = Provider.objects.filter(id__in=providers_ids)
        else:
            providers = Provider.objects.filter(units__in=reservation_services.values_list('unit', flat=True)).distinct()
        for provider in providers:
            reservation_services_filtered = reservation_services.filter(unit__provider=provider)
            units_group = self.units_agroups_report(reservation_services_filtered)
            reservation_services_groups = []
            for unit_value in units_group:
                reservation_services_group_query = reservation_services_filtered.filter(unit_value=unit_value)
                if reservation_services_group_query.exists():
                    reservation_services_group = {
                        'reservation_services':reservation_services_group_query,
                    }
                    reservation_services_group.update(reservation_services_group_query.paxAggregate())
                    reservation_services_groups.append(reservation_services_group)
            if provider.email != "":
                if self.type == "pdf":
                    result, filename = self.report_in_pdf(reservation_services_groups,reservation_services_filtered)
                else:
                    result, filename = self.report_in_excel(reservation_services_groups)

                email = views.EmailsViewSet.prepare_html_email_from_template_pdf_attach("emails/report-provider/template.html",
                    {
                        'provider' : provider,
                        'date' : self.date,
                    },
                    [provider.email],
                    filename,
                    result
                )
                email.send()
    
    def report_in_pdf(self,reservation_services_groups,reservation_services):
        from GeneralApp.models import Hotel
        template = get_template('operation_report.html')
        context = {
            'host'                              :   self.request.get_host(),
            'reservation_services_groups'       :   reservation_services_groups,
            'date'                              :   self.date,
            'today'                             :   datetime.now(),
            'user'                              :   self.request.user,
            'property'                          :   property,
            'unit'                              :   self.unit,
        }
        context.update(reservation_services.paxAggregate())
        html  = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            filename = "{}-operation-report-{}-{}.pdf".format(str(self.date),self.property.code,str(datetime.now()))
            return result, filename
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
    
    def report_in_excel(self,reservation_services_groups):
        tranfer_type_options = {
            'DEPARTURES':'S',
            'ARRIVALS':'L',
            'INTERHOTEL':'I',
        }
        output = BytesIO()
        filename = '{}-operation-report-{}-{}.xlsx'.format(str(self.date),self.property.code,str(datetime.now()))
        workbook = xlsxwriter.Workbook(output)
        merge_format = workbook.add_format({
            "bold": 1,
            "align": "center",
            "valign": "vcenter",
        })
        header_format = workbook.add_format({
            'bold': True,
            'top':1,
            'bottom':1,
            "align": "center",
            "valign": "vcenter",
            'font_size':15,
        })
        unit_format = workbook.add_format({
            'bold': True,
            "align": "left",
            "valign": "vcenter",
            'font_size':12,
        })
        worksheet = workbook.add_worksheet()
        worksheet.merge_range("A1:L1", "OPERACIÓN DE TRASLADOS - {}".format(self.property.name), merge_format)
        worksheet.merge_range("A2:L2", "DEL {}".format(self.date.strftime("%d/%m/%Y")), merge_format)
        #worksheet.merge_range("A3:L3", "BOULEVARD KUKULKAN KM 14.5 MANZANA 53 LOTE 40 ZONA HOTELERA", merge_format)
        worksheet.write("A3", "Pasajero", header_format)
        worksheet.set_column('A:A', 25)
        worksheet.write("B3", "Referencia", header_format)
        worksheet.set_column('B:B', 15)
        worksheet.write("C3", "Ad", header_format)
        worksheet.write("D3", "Ni", header_format)
        worksheet.set_column('C:D', 5)
        worksheet.write("E3", "Vuelo", header_format)
        worksheet.set_column('E:E', 20)
        worksheet.write("F3", "Hora", header_format)
        worksheet.set_column('F:F', 7)
        worksheet.write("G3", "Pickup", header_format)
        worksheet.set_column('G:G', 15)
        worksheet.write("H3", "Hotel", header_format)
        worksheet.set_column('H:H', 30)
        worksheet.write("I3", "Destino", header_format)
        worksheet.set_column('I:I', 30)
        worksheet.write("J3", "HAB", header_format)
        worksheet.set_column('J:J', 7)
        worksheet.write("K3", "Notas", header_format)
        worksheet.set_column('K:K', 30)
        worksheet.write("L3", "Tipo Servicio", header_format)
        worksheet.set_column('L:L', 20)
        worksheet.write("M3", "UNIDAD", header_format)
        worksheet.set_column('M:M', 15)
        row = 3
        for reservation_services_group in reservation_services_groups:
            for reservation_service in reservation_services_group['reservation_services'].all():
                worksheet.write(row,0, reservation_service.reservation.pax)
                worksheet.write(row,1, '{0:03d}'.format(reservation_service.reservation.id))
                worksheet.write(row,2, reservation_service.adults)
                worksheet.write(row,3, reservation_service.childs)
                worksheet.write(row,4, reservation_service.flight_code if reservation_service.flight_code!="" else "TBA")
                if reservation_service.real_flight_time is not None and reservation_service.real_flight_time != "":
                    worksheet.write(row,5, reservation_service.real_flight_time.strftime("%H:%M"))
                elif reservation_service.flight is not None:
                    flight_time = getattr(reservation_service.flight, reservation_service.flight_field)
                    if flight_time is not None:
                        worksheet.write(row,5, flight_time.strftime("%H:%M"))
                    else:
                        worksheet.write(row,5, "__:__")
                else:
                    worksheet.write(row,5, "__:__")
                cell_format = workbook.add_format()
                cell_format.set_bold()
                worksheet.write(row, 6, reservation_service.pup.strftime("%H:%M") if reservation_service.pup is not None else "", cell_format)
                if reservation_service.origin is not None:
                    worksheet.write(row,7, reservation_service.origin.name)
                elif reservation_service.destination is not None:
                    worksheet.write(row,7, reservation_service.destination.name)
                if reservation_service.transfer_type == "INTERHOTEL" and reservation_service.destination is not None:
                    worksheet.write(row,8, reservation_service.destination.name)
                
                worksheet.write(row,9, reservation_service.room if reservation_service.room is not None else "")
                worksheet.write(row,10, '{} | {}'.format(reservation_service.reservation_comment,reservation_service.comments))
                worksheet.write(row,11,'{}-{}'.format(tranfer_type_options[reservation_service.transfer_type],reservation_service.service.code))
                worksheet.write(row,12, reservation_service.unit_name if reservation_service.unit_name is not None else "")
                row += 1
            """ worksheet.write(row,11, reservation_services_group['pax_total'], unit_format)
            row += 1 """
        workbook.close()
        output.seek(0)
        return output, filename
    
class OperationListController():
    def __init__(self, request, start_date, due_date, property, type, format, sale_types, hotels, services, operation_types, reservation_confirm):
        self.request = request
        self.start_date = start_date
        self.due_date = due_date
        self.property = property
        self.type = type
        self.format = format
        self.sale_types = sale_types
        self.operation_types = operation_types
        self.hotels = hotels
        self.services = services
        self.reservation_confirm = reservation_confirm

    def title(self):
        transfers = dict(TRANSFER_TYPE_CHOICES)
        if self.type != "ALL":
            return "({} {})".format(transfers[self.type], self.reservation_confirm)
        return "({} {})".format("Todas", self.reservation_confirm)

    def filters(self):
        reservation_services = models.ReservationService.objects.tableOptimization(
            ).operationFilterDateRange(
                self.start_date,self.due_date,self.property
            ).commentsAnnotate(
            ).exclude(
                reservation__status="CANCEL"
            ).order_by_time()
        if self.type != "ALL":
            reservation_services = reservation_services.filter(transfer_type=self.type)
        if len(self.sale_types) > 0:
            reservation_services = reservation_services.operationFilterBySaleTypes(self.sale_types)
        if len(self.operation_types) > 0:
            reservation_services = reservation_services.operationFilterByOperationTypes(self.operation_types)
        if len(self.hotels) > 0:
            reservation_services = reservation_services.operationFilterByHotels(self.hotels)
        if len(self.services) > 0:
            reservation_services = reservation_services.operationFilterByServices(self.services)
        if self.type == "DEPARTURES":
            if self.reservation_confirm == "CONF":
                reservation_services = reservation_services.filter(confirmation=True)
            if self.reservation_confirm == "NO CONF":
                reservation_services = reservation_services.filter(confirmation=False)
        return reservation_services
    
    def get_context(self):
        reservation_services = self.filters()
        dates = reservation_services.operationGroupByDay()
        reservation_services_groups = []
        for date in dates:
            reservation_services_group = {}
            reservation_services_query = reservation_services.filter(operation_date=date).order_by('pup')
            print(reservation_services_query)
            reservation_services_group['reservation_services'] = reservation_services_query
            reservation_services_group['date'] = date
            reservation_services_group.update(reservation_services_query.paxTotal())
            reservation_services_groups.append(reservation_services_group)
        context = {
            'host'                              :   self.request.get_host(),
            'reservation_services_groups'       :   reservation_services_groups,
            'start_date'                        :   self.start_date,
            'due_date'                          :   self.due_date,
            'property'                          :   self.property,
            'today'                             :   datetime.now(),
            'user'                              :   self.request.user,
            'title'                             :   self.title()
        }
        context.update(reservation_services.paxAggregate())
        if self.format == "pdf":
            return self.report_in_pdf(context)
        elif self.format == "excel":
            return self.report_in_excel(context)
        elif self.format == "word":
            return self.report_in_word(context)
        else:
            raise CustomValidation("Formato invalido", 'report', status.HTTP_400_BAD_REQUEST)

    def report_in_pdf(self,context):
        if self.type == "INTERHOTEL":
            template = get_template('operation_report_list_interhotel.html')
        else:
            template = get_template('operation_report_list_arrivals_departures.html')
        html  = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            filename = "{}-{}-operation-report-list-{}-{}.pdf".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
            response = HttpResponse(
                result.getvalue(), 
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = "inline; filename={}".format(filename)
            return response
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
            
    def report_in_excel(self,context):
        from django.http import FileResponse
        if self.type == "INTERHOTEL":
            result, filename = self.excel_interhotel(context)
        else:
            result, filename = self.excel_arrivals_departures(context)
        response = HttpResponse(
            result.getvalue(),
            content_type='application/ms-excel;charset=utf-8'
        )
        response['Content-Disposition'] = "attachment; filename={}".format(filename)
        response["Content-Encoding"] = "UTF-8"
        return response
        
    def excel_interhotel(self, context):
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output)
        merge_format = workbook.add_format({
            "bold": 1,
            "align": "center",
            "valign": "vcenter",
        })
        header_format = workbook.add_format({
            'bold': True,
            'top':1,
            'bottom':1,
            "align": "center",
            "valign": "vcenter",
            'font_size':15,
        })
        unit_format = workbook.add_format({
            'bold': True,
            "align": "left",
            "valign": "vcenter",
            'font_size':12,
            'top':1,
        })
        
        worksheet = workbook.add_worksheet()
        worksheet.merge_range("A1:K1", "ON TOUR BY RCD - {} {}".format(self.property.name,context['title']), merge_format)
        worksheet.merge_range("A2:K2", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
        worksheet.merge_range("A3:K3", "{} {}".format(context["user"].username,context['today'].strftime("%d/%m/%Y %H:%M:%S")), merge_format)
        worksheet.write("A4", "Fecha", header_format)
        worksheet.set_column('A:A', 15)
        worksheet.write("B4", "Pasajero", header_format)
        worksheet.set_column('B:B', 25)
        worksheet.write("C4", "CveRes", header_format)
        worksheet.set_column('C:C', 15)
        worksheet.write("D4", "Ad", header_format)
        worksheet.write("E4", "Ni", header_format)
        worksheet.set_column('D:E', 5)
        worksheet.write("F4", "Tipo de venta", header_format)
        worksheet.set_column('F:F', 20)
        worksheet.write("G4", "Hotel Origen/Destino", header_format)
        worksheet.set_column('G:G', 40)
        worksheet.write("H4", "CveOpera", header_format)
        worksheet.set_column('H:H', 20)
        worksheet.write("I4", "Notas", header_format)
        worksheet.set_column('I:I', 30)
        worksheet.write("J4", "TipoOpera", header_format)
        worksheet.set_column('J:J', 20)
        worksheet.write("K4", "Tipo Serv", header_format)
        worksheet.set_column('K:K', 20)
        worksheet.write("M4", "Unidad", header_format)
        worksheet.set_column('M:M', 20)
        row = 4
        for reservation_services_group in context['reservation_services_groups']:
            for reservation_service in reservation_services_group['reservation_services'].all():
                worksheet.write(row,0, reservation_service.operation_date.strftime("%d/%m/%Y"))
                worksheet.write(row,1, reservation_service.reservation.pax)
                worksheet.write(row,2, '{0:03d}'.format(reservation_service.reservation.id))
                worksheet.write(row,3, reservation_service.adults)
                worksheet.write(row,4, reservation_service.childs)
                worksheet.write(row,5, reservation_service.reservation.sale_type.name)
                worksheet.write(row,6, "{} \n {}".format(reservation_service.origin.name,reservation_service.destination.name))
                worksheet.write(row,7, "{}".format(reservation_service.reservation.opera_code))
                worksheet.write(row,8, '{} // {}'.format(reservation_service.reservation_comment,reservation_service.comments))
                worksheet.write(row,9, reservation_service.operation_type.name)
                worksheet.write(row,10, reservation_service.service.code)
                worksheet.write(row,11, reservation_service.unit_name if reservation_service.unit_pk is not None else "")
                row += 1
            worksheet.merge_range("A{}:C{}".format(row+1,row+1), "TOTAL DIA {}".format(reservation_services_group['date'].strftime("%d/%m/%Y")), unit_format)
            worksheet.write(row,3, reservation_services_group['adults_total'],unit_format)
            worksheet.write(row,4, reservation_services_group['childs_total'],unit_format)
            row += 1
        worksheet.write(row,10, "TOTAL {}".format(context['pax_total']), unit_format)
        row += 1
        workbook.close()
        output.seek(0)
        filename = "{}-{}-operation-report-list-{}-{}.xlsx".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        return output, filename
    
    def excel_arrivals_departures(self, context):
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output)
        merge_format = workbook.add_format({
            "bold": 1,
            "align": "center",
            "valign": "vcenter",
        })
        header_format = workbook.add_format({
            'bold': True,
            'top':1,
            'bottom':1,
            "align": "center",
            "valign": "vcenter",
            'font_size':15,
        })
        unit_format = workbook.add_format({
            'bold': True,
            "align": "left",
            "valign": "vcenter",
            'font_size':12,
            'top':1,
        })
        
        worksheet = workbook.add_worksheet()
        worksheet.merge_range("A1:M1", "ON TOUR BY RCD - {} {}".format(self.property.name,context['title']), merge_format)
        worksheet.merge_range("A2:M2", "{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")), merge_format)
        worksheet.merge_range("A3:M3", "{} {}".format(context["user"].username,context['today'].strftime("%d/%m/%Y %H:%M:%S")), merge_format)
        worksheet.write("A4", "Fecha", header_format)
        worksheet.set_column('A:A', 15)
        worksheet.write("B4", "Pasajero", header_format)
        worksheet.set_column('B:B', 25)
        worksheet.write("C4", "CveRes", header_format)
        worksheet.set_column('C:C', 15)
        worksheet.write("D4", "Ad", header_format)
        worksheet.write("E4", "Ni", header_format)
        worksheet.set_column('D:E', 5)
        worksheet.write("F4", "Vuelo", header_format)
        worksheet.set_column('F:F', 20)
        worksheet.write("G4", "Hora", header_format)
        worksheet.set_column('G:G', 7)
        worksheet.write("H4", "Tipo de venta", header_format)
        worksheet.set_column('H:H', 20)
        worksheet.write("I4", "Hotel", header_format)
        worksheet.set_column('I:I', 30)
        worksheet.write("J4", "CveOpera", header_format)
        worksheet.set_column('J:J', 20)
        worksheet.write("K4", "Notas", header_format)
        worksheet.set_column('K:K', 30)
        worksheet.write("L4", "TipoOpera", header_format)
        worksheet.set_column('L:L', 20)
        worksheet.write("M4", "Tipo Serv", header_format)
        worksheet.set_column('M:M', 20)
        worksheet.write("N4", "Unidad", header_format)
        worksheet.set_column('N:N', 20)
        row = 4
        for reservation_services_group in context['reservation_services_groups']:
            for reservation_service in reservation_services_group['reservation_services'].all():
                worksheet.write(row,0, reservation_service.operation_date.strftime("%d/%m/%Y"))
                worksheet.write(row,1, reservation_service.reservation.pax)
                worksheet.write(row,2, '{0:03d}'.format(reservation_service.reservation.id))
                worksheet.write(row,3, reservation_service.adults)
                worksheet.write(row,4, reservation_service.childs)
                worksheet.write(row,5, reservation_service.flight_code if reservation_service.flight_code!="" else "TBA")
                if reservation_service.flight_time:
                    worksheet.write(row,6, reservation_service.flight_time.strftime("%H:%M"))
                else:
                    worksheet.write(row,6, "__:__")
                worksheet.write(row,7, reservation_service.reservation.sale_type.name)
                if reservation_service.destination is not None:
                    worksheet.write(row,8, reservation_service.destination.name)
                    worksheet.write(row,9, reservation_service.reservation.opera_code)
                elif reservation_service.origin is not None:
                    worksheet.write(row,8, reservation_service.origin.name)
                    worksheet.write(row,9, reservation_service.reservation.opera_code)
                
                worksheet.write(row,10, '{} // {}'.format(reservation_service.reservation_comment,reservation_service.comments))
                worksheet.write(row,11, reservation_service.operation_type.name)
                worksheet.write(row,12, reservation_service.service.code)
                worksheet.write(row,13, reservation_service.unit_name if reservation_service.unit_pk is not None else "")
                row += 1
            worksheet.merge_range("A{}:C{}".format(row+1,row+1), "TOTAL DIA {}".format(reservation_services_group['date'].strftime("%d/%m/%Y")), unit_format)
            worksheet.write(row,3, reservation_services_group['adults_total'],unit_format)
            worksheet.write(row,4, reservation_services_group['childs_total'],unit_format)
            row += 1
        worksheet.write(row,12, "TOTAL {}".format(context['pax_total']), unit_format)
        row += 1
        workbook.close()
        output.seek(0)
        filename = "{}-{}-operation-report-list-{}-{}.xlsx".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        return output, filename
    
    def report_in_word(self,context):
        if self.type == "INTERHOTEL":
            result, filename = self.word_interhotel(context)
        else:
            result, filename = self.word_arrivals_departures(context)
        response = HttpResponse(
            result.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response['Content-Disposition'] = "attachment; filename={}".format(filename)
        response["Content-Encoding"] = "UTF-8"
        return response
        
    def word_interhotel(self, context):
        output = BytesIO()
        document = Document()

        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.orientation, section.page_width, section.page_height

        style = document.styles['Normal']
        font = style.font
        font.name ='Arial'
        font.size = Pt(10)

        document.add_heading("ON TOUR BY RCD - {} {}".format(self.property.name,context['title']))
        document.add_paragraph("{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")))
        document.add_paragraph("{} {}".format(context["user"].username,context['today'].strftime("%d/%m/%Y %H:%M:%S")))

        table = document.add_table(rows=1,cols=12,style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Fecha'
        hdr_cells[1].text = 'Pasajero'
        hdr_cells[2].text = 'CveRes'
        hdr_cells[3].text = 'Ad'
        hdr_cells[4].text = 'Ni'
        hdr_cells[5].text = 'Tipo de venta'
        hdr_cells[6].text = 'Hotel Origen/Destino'
        hdr_cells[7].text = 'CveOpera'
        hdr_cells[8].text = 'Notas'
        hdr_cells[9].text = 'TipoOpera'
        hdr_cells[10].text = 'Tipo Serv'
        hdr_cells[11].text = 'Unidad'

        for cell in table.rows[0].cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.underline = True
        row = 0
        for reservation_services_group in context['reservation_services_groups']:
            for reservation_service in reservation_services_group['reservation_services'].all():
                row_cells = table.add_row().cells
                row += 1
                row_cells[0].text = reservation_service.operation_date.strftime("%d/%m/%Y")
                row_cells[1].text = reservation_service.reservation.pax
                row_cells[2].text = '{0:03d}'.format(reservation_service.reservation.id)
                row_cells[3].text = str(reservation_service.adults)
                row_cells[4].text = str(reservation_service.childs)
                row_cells[5].text = reservation_service.reservation.sale_type.name
                row_cells[6].text = "{}\n{}".format(reservation_service.origin.name,reservation_service.destination.name)
                row_cells[7].text = "{}".format(reservation_service.reservation.opera_code)
                row_cells[8].text = '{} // {}'.format(reservation_service.reservation_comment,reservation_service.comments)
                row_cells[9].text = reservation_service.operation_type.name
                row_cells[10].text = reservation_service.service.code
                row_cells[11].text = reservation_service.unit_name if reservation_service.unit_pk is not None else ""
            
            row_cells = table.add_row().cells
            row += 1
            row_cells[0].text = "TOTAL DIA {}".format(reservation_services_group['date'].strftime("%d/%m/%Y"))
            row_cells[0].style = ""
            table.cell(row, 0).merge(table.cell(row, 2))

            row_cells[3].text = str(reservation_services_group['adults_total'])
            row_cells[4].text = str(reservation_services_group['childs_total'])
        row_cells = table.add_row().cells
        row += 1
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)
        row_cells[10].text = "TOTAL {}".format(context['pax_total'])
        filename = "{}-{}-operation-report-list-{}-{}.docx".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        document.save(output)
        output.seek(0)
        return output, filename
    
    def word_arrivals_departures(self, context):
        output = BytesIO()
        document = Document()

        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.orientation, section.page_width, section.page_height

        style = document.styles['Normal']
        font = style.font
        font.name ='Arial'
        font.size = Pt(10)

        document.add_heading("ON TOUR BY RCD - {} {}".format(self.property.name,context['title']))
        document.add_paragraph("{} - {}".format(context['start_date'].strftime("%d/%m/%Y"),context['due_date'].strftime("%d/%m/%Y")))
        document.add_paragraph("{} {}".format(context["user"].username,context['today'].strftime("%d/%m/%Y %H:%M:%S")))

        table = document.add_table(rows=1,cols=14,style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Fecha'
        hdr_cells[1].text = 'Pasajero'
        hdr_cells[2].text = 'CveRes'
        hdr_cells[3].text = 'Ad'
        hdr_cells[4].text = 'Ni'
        hdr_cells[5].text = 'Vuelo'
        hdr_cells[6].text = 'Hora'
        hdr_cells[7].text = 'Tipo de venta'
        hdr_cells[8].text = 'Hotel'
        hdr_cells[9].text = 'CveOpera'
        hdr_cells[10].text = 'Notas'
        hdr_cells[11].text = 'TipoOpera'
        hdr_cells[12].text = 'Tipo Serv'
        hdr_cells[13].text = 'Unidad'
        for cell in table.rows[0].cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.underline = True
        row = 0
        for reservation_services_group in context['reservation_services_groups']:
            for reservation_service in reservation_services_group['reservation_services'].all():
                row_cells = table.add_row().cells
                row += 1
                row_cells[0].text = reservation_service.operation_date.strftime("%d/%m/%Y")
                row_cells[1].text = reservation_service.reservation.pax
                row_cells[2].text = '{0:03d}'.format(reservation_service.reservation.id)
                row_cells[3].text = str(reservation_service.adults)
                row_cells[4].text = str(reservation_service.childs)
                row_cells[5].text = reservation_service.flight_code if reservation_service.flight_code!="" else "TBA"
                if reservation_service.flight_time is not None:
                    row_cells[6].text = reservation_service.flight_time.strftime("%H:%M")
                else:
                    row_cells[6].text = "__:__"
                row_cells[7].text = reservation_service.reservation.sale_type.name
                if reservation_service.destination is not None:
                    row_cells[8].text = reservation_service.destination.name
                    row_cells[9].text = reservation_service.reservation.opera_code
                elif reservation_service.origin is not None:
                    row_cells[8].text = reservation_service.origin.name
                    row_cells[9].text = reservation_service.reservation.opera_code
                row_cells[10].text = '{} // {}'.format(reservation_service.reservation_comment,reservation_service.comments)
                row_cells[11].text = reservation_service.operation_type.name
                row_cells[12].text = reservation_service.service.code
                row_cells[13].text = reservation_service.unit_name if reservation_service.unit_pk is not None else ""
            
            row_cells = table.add_row().cells
            row += 1
            row_cells[0].text = "TOTAL DIA {}".format(reservation_services_group['date'].strftime("%d/%m/%Y"))
            row_cells[0].style = ""
            table.cell(row, 0).merge(table.cell(row, 2))

            row_cells[3].text = str(reservation_services_group['adults_total'])
            row_cells[4].text = str(reservation_services_group['childs_total'])
        row_cells = table.add_row().cells
        row += 1
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)
        row_cells[12].text = "TOTAL {}".format(context['pax_total'])
        filename = "{}-{}-operation-report-list-{}-{}.docx".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        document.save(output)
        output.seek(0)
        return output, filename

class OperationCouponController():
    def __init__(self, request, start_date, due_date, property, type, sale_types, hotels, services, operation_types):
        self.request = request
        self.start_date = start_date
        self.due_date = due_date
        self.property = property
        self.type = type
        self.sale_types = sale_types
        self.operation_types = operation_types
        self.hotels = hotels
        self.services = services

    def filters(self):
        reservation_services = models.ReservationService.objects.operationFilterDateRange(
                self.start_date,self.due_date,self.property
            ).commentsAnnotate(
            ).tableOptimization(
            ).exclude(
                reservation__status="CANCEL"
            )
        if self.type != "ALL":
            reservation_services = reservation_services.filter(transfer_type=self.type)
        if len(self.sale_types) > 0:
            reservation_services = reservation_services.operationFilterBySaleTypes(self.sale_types)
        if len(self.operation_types) > 0:
            reservation_services = reservation_services.operationFilterByOperationTypes(self.operation_types)
        if len(self.hotels) > 0:
            reservation_services = reservation_services.operationFilterByHotels(self.hotels)
        if len(self.services) > 0:
            reservation_services = reservation_services.operationFilterByServices(self.services)
        return reservation_services.order_by('operation_date','pup')
    
    def get_context(self):
        reservation_services = self.filters()
        reservation_services_data = []
        i = 1
        for reservation_service in reservation_services:
            reservation_service_data = serializers.ReservationServiceOperationCouponSerializer(reservation_service).data
            reservation_service_data['cveres'] = reservation_service.reservation.id
            reservation_service_data['operation_date'] = reservation_service.operation_date.strftime("%d/%m/%Y")
            reservation_service_data['pax_total'] = int(reservation_service.adults + reservation_service.childs)
            if reservation_service.flight_time is not None:
                reservation_service_data['flight_time'] = reservation_service.flight_time.strftime("%H:%M")
            else:
                reservation_service_data['flight_time'] = "__:__"
            reservation_service_data['departure_date'] = reservation_service.departure_date.strftime("%d-%b-%Y") if reservation_service.departure_date is not None else ""
            reservation_service_data['pup'] = reservation_service.pup.strftime("%H:%M") if reservation_service.pup is not None else ""
            service_type = "Serv.Colectivo" if reservation_service.service.is_colective else "Serv.Privado"
            if reservation_service.transfer_type == "INTERHOTEL":
                service_type += "-Interhotel"
            reservation_type = ""
            if reservation_service.reservation_type == 'RT':
                reservation_type = "ROUNDTRIP"
            else:
                reservation_type = "ONE WAY"
            if reservation_service.service.is_colective is True:
                reservation_service_data['service_detail'] = "{} - {} ".format(service_type,reservation_type)
            else:
                reservation_service_data['service_detail'] = "{} - {} ({})".format(service_type,reservation_type,reservation_service.service.unit.name if reservation_service.service.unit is not None else "")
            if reservation_service.transfer_type == "ARRIVALS":            
                departure = reservation_service.reservation.reservation_services.filter(transfer_type='DEPARTURES').order_by("-date").first()
                if departure is not None:
                    reservation_service_data['departure_hotel'] = departure.origin.name
                    reservation_service_data['flight_out_code'] = departure.flight_code
                    flight_out_time = None
                    departure_pup = None
                    if departure.real_flight_time is not None and departure.real_flight_time != "":
                        reservation_service_data['flight_out_time'] = departure.real_flight_time.strftime("%H:%M")
                        flight_out_time = departure.real_flight_time
                    elif departure.flight is not None:
                        reservation_service_data['flight_out_time'] = getattr(departure.flight, departure.flight_field).strftime("%H:%M")
                        flight_out_time = getattr(departure.flight, departure.flight_field)
                    if departure.real_pick_up_time is not None and departure.real_pick_up_time != "":
                        reservation_service_data['departure_pup'] = departure.real_pick_up_time.strftime("%H:%M")
                        departure_pup = departure.real_pick_up_time
                    elif departure.pick_up_time is not None:
                        reservation_service_data['departure_pup'] = departure.pick_up_time.time.strftime("%H:%M")
                        departure_pup = departure.pick_up_time.time
                    if flight_out_time is not None and departure_pup is not None:
                        if departure_pup > flight_out_time:
                            departure_date = reservation_service.departure_date - timedelta(days=1)
                            reservation_service_data['departure_date'] = departure_date.strftime("%d-%b-%Y")
                else:
                    reservation_service_data['flight_out_code'] = ""
                    reservation_service_data['flight_out_time'] = ""
                    reservation_service_data['departure_pup'] = ""
                    reservation_service_data['departure_hotel'] = ""
                
            reservation_service_data['break_page'] = i%3 == 0
            i += 1
            reservation_services_data.append(reservation_service_data)
        context = {
            'host'                              :   self.request.get_host(),
            'environment'                       :   getattr(serverconfig,"environment","http"),
            'reservation_services'              :   reservation_services_data,
            'start_date'                        :   self.start_date,
            'due_date'                          :   self.due_date,
            'property'                          :   self.property,
            'today'                             :   datetime.now(),
            'user'                              :   self.request.user,
        }
        template = get_template('operation_report_coupons.html')
        html  = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            filename = "{}-operation-report-{}-{}.pdf".format(str(self.start_date),self.property.code,str(datetime.now()))
            response = HttpResponse(result.getvalue(), content_type='application/pdf')
            content = "inline; filename={}".format(filename)
            response['Content-Disposition'] = content
            return response
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
        """ pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            filename = "{}-{}-operation-report-coupons-{}-{}.pdf".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
            response = HttpResponse(
                result.getvalue(), 
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = "inline; filename={}".format(filename)
            return response
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST) """

class OperationSummaryController():
    def __init__(self, request, start_date, due_date, year, property, type, print_total, format, sort_by, sale_types, hotels):
        self.request = request
        self.start_date = start_date
        self.due_date = due_date
        self.year = year
        self.property = property
        self.type = type
        self.format = format
        self.sale_types = sale_types
        self.sort_by = sort_by
        self.hotels = hotels
        self.print_total = print_total

    def title(self):
        transfers = dict(TRANSFER_TYPE_CHOICES)
        return "{}".format(transfers[self.type])
    
    def title_hotels(self):
        from GeneralApp.models import Hotel
        if len(self.hotels) > 0:
            return "HOTELES: [{}]".format(','.join(Hotel.objects.filter(id__in=self.hotels).values_list("name", flat=True)))
        else:
            return "HOTELES: [TODOS]"
    
    def title_sale_types(self):
        from GeneralApp.models import SaleType
        if len(self.sale_types) > 0:
            return "TIPO DE VENTAS: [{}]".format(','.join(SaleType.objects.filter(id__in=self.hotels).values_list("name", flat=True)))
        else:
            return "TIPO DE VENTAS: [TODOS]"

    def filters(self):
        if self.year is not None:
            reservation_services = models.ReservationService.objects.tableOptimization(
                ).operationFilterDateYear(
                    self.year,self.property
                ).filter(
                    transfer_type=self.type
                ).exclude(
                    reservation__status="CANCEL"
                )
        else:
            delta = self.due_date - self.start_date
            if delta.days > 30:
                raise CustomValidation("El rango de fechas es demasiado grande", 'date', status.HTTP_400_BAD_REQUEST)
            reservation_services = models.ReservationService.objects.tableOptimization(
                ).operationFilterDateRange(
                    self.start_date,self.due_date,self.property
                ).filter(
                    transfer_type=self.type
                ).exclude(
                    reservation__status="CANCEL"
                )
        if len(self.sale_types) > 0:
            reservation_services = reservation_services.operationFilterBySaleTypes(self.sale_types)
        if len(self.hotels) > 0:
            reservation_services = reservation_services.operationFilterByHotels(self.hotels)
        return reservation_services
    
    def get_context(self):
        reservation_services = self.filters()
        sort_by = []
        if self.sort_by == "sales_type":
            sort_by = reservation_services.operationGroupBySaleType()
        if self.sort_by == "hotel":
            if self.type == 'DEPARTURES':
                sort_by = reservation_services.operationGroupByHotelOrigin()
            if self.type == 'ARRIVALS':
                sort_by = reservation_services.operationGroupByHotelDestination()

        if self.year is not None:
            context = self.report_by_year(reservation_services,sort_by)
        else:
            context = self.report_by_date(reservation_services,sort_by)
        if self.format == "pdf":
            return self.report_in_pdf(context)
        elif self.format == "excel":
            return self.report_in_excel(context)
        elif self.format == "word":
            return self.report_in_word(context)
        else:
            return Response(context)
        
    def report_by_year(self,reservation_services,sort_by):
        reservation_services_groups = []
        total_by_dates = []
        for sort in sort_by:
            reservation_services_group = {}
            reservation_services_group['name'] = sort['name']
            reservation_services_group['dates'] = []
            if self.sort_by == "sales_type":
                reservation_services_query = reservation_services.filter(reservation__sale_type__id=sort['id_'])
            else:
                if self.type == 'DEPARTURES':
                    reservation_services_query = reservation_services.filter(origin__id=sort['id_'])
                if self.type == 'ARRIVALS':
                    reservation_services_query = reservation_services.filter(destination__id=sort['id_']) 
            for month in [1,2,3,4,5,6,7,8,9,10,11,12]:
                total = reservation_services_query.filter(date__month=month).paxTotal()
                reservation_services_group['dates'].append(total)
                item = next((item for item in total_by_dates if item["date"] == month), None)
                if item is not None:
                    item["adults_total"] += total["adults_total"] if total["adults_total"] is not None else 0
                    item["childs_total"] += total["childs_total"] if total["childs_total"] is not None else 0
                else:
                    total_by_dates.append({
                        'date': month,
                        'adults_total': total["adults_total"] if total["adults_total"] is not None else 0,
                        'childs_total': total["childs_total"] if total["childs_total"] is not None else 0,
                    })
            reservation_services_group.update(reservation_services_query.paxTotal())
            reservation_services_groups.append(reservation_services_group)
        context = {
            'host'                              :   self.request.get_host(),
            'reservation_services_groups'       :   reservation_services_groups,
            'year'                              :   self.year,
            'type'                              :   self.type,
            'sort_by'                           :   self.sort_by,
            'print_total'                       :   self.print_total,
            'property'                          :   self.property,
            'today'                             :   datetime.now(),
            'user'                              :   self.request.user,
            'title'                             :   self.title(),
            'title_hotels'                      :   self.title_hotels(),
            'title_sale_types'                  :   self.title_sale_types(),
            'total_by_dates'                    :   total_by_dates 
        }
        context.update(reservation_services.paxTotal())
        return context

    def report_by_date(self,reservation_services,sort_by):
        dates = reservation_services.operationGroupByDay()
        total_by_dates = []
        reservation_services_groups = []
        class_table = True
        for sort in sort_by:
            reservation_services_group = {}
            reservation_services_group['name'] = sort['name']
            reservation_services_group['class_table'] = "gray" if class_table is True else ""
            class_table = not class_table
            reservation_services_group['dates'] = []
            if self.sort_by == "sales_type":
                reservation_services_query = reservation_services.filter(reservation__sale_type__id=sort['id_'])
            else:
                if self.type == 'DEPARTURES':
                    reservation_services_query = reservation_services.filter(origin__id=sort['id_'])
                if self.type == 'ARRIVALS':
                    reservation_services_query = reservation_services.filter(destination__id=sort['id_'])
            for date in daterange(self.start_date, self.due_date):
                total = reservation_services_query.filter(operation_date=date).paxTotal()
                reservation_services_group['dates'].append(total)
                item = next((item for item in total_by_dates if item["date"] == date), None)
                if item is not None:
                    item["adults_total"] += total["adults_total"] if total["adults_total"] is not None else 0
                    item["childs_total"] += total["childs_total"] if total["childs_total"] is not None else 0
                else:
                    total_by_dates.append({
                        'date': date,
                        'adults_total': total["adults_total"] if total["adults_total"] is not None else 0,
                        'childs_total': total["childs_total"] if total["childs_total"] is not None else 0,
                    })
            reservation_services_group.update(reservation_services_query.paxTotal())
            reservation_services_groups.append(reservation_services_group)
        context = {
            'host'                              :   self.request.get_host(),
            'reservation_services_groups'       :   reservation_services_groups,
            'start_date'                        :   self.start_date,
            'due_date'                          :   self.due_date,
            'type'                              :   self.type,
            'sort_by'                           :   self.sort_by,
            'print_total'                       :   self.print_total,
            'property'                          :   self.property,
            'today'                             :   datetime.now(),
            'user'                              :   self.request.user,
            'title'                             :   self.title(),
            'title_hotels'                      :   self.title_hotels(),
            'title_sale_types'                  :   self.title_sale_types(),
            'total_by_dates'                    :   total_by_dates
        }
        context.update(reservation_services.paxTotal())
        return context

    def report_in_pdf(self,context):
        if self.year is not None:
            template = get_template('operation_report_summary_by_year.html')
            filename = "{}-operation-report-summary-{}-{}.pdf".format(str(self.year),self.property.code,str(datetime.now()))
        else:
            template = get_template('operation_report_summary_by_dates.html')
            filename = "{}-{}-operation-report-summary-{}-{}.pdf".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        html  = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            
            response = HttpResponse(
                result.getvalue(), 
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = "inline; filename={}".format(filename)
            return response
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)
  
    def report_in_excel(self, context):
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output)
        merge_format = workbook.add_format({
            "bold": 1,
            "align": "center",
            "valign": "vcenter",
        })
        merge_format_left = workbook.add_format({
            "bold": 1,
            "align": "left",
            "valign": "vcenter",
        })
        header_format = workbook.add_format({
            'bold': True,
            'top':1,
            'bottom':1,
            "align": "center",
            "valign": "vcenter",
            'font_size':15,
        })
        unit_format = workbook.add_format({
            'bold': True,
            "align": "left",
            "valign": "vcenter",
            'font_size':12,
            'top':1,
        })
        total_format = workbook.add_format({
            'bold': True,
            "align": "left",
            "valign": "vcenter",
            'font_size':8.5,
            'bottom':1,
        })
        total_format_ = workbook.add_format({
            'bold': True,
            "valign": "vcenter",
            'text_wrap': True,
            "align": "center",
            'font_size':12,
            'font_color':"#008be8"
        })
        worksheet = workbook.add_worksheet('Resumen')
        months = []
        if self.year is not None:
            months = ['ENE','FEB','MAR','ABR','MAY','JUN','JUL','AGO','SEP','OCT','NOV','DIC']
            worksheet.merge_range("A1:F3", "ON TOUR BY RCD - {}".format(self.property.name[4:]), merge_format)
            type_title = "POR TIPO DE VENTA" if context['sort_by'] == "sales_type" else "POR HOTEL"
            worksheet.merge_range(
                "G1:M3",
                "RESUMEN DE {} {} PARA EL AÑO {} \n Reporte impreso por: {} el {}".format(
                    context['title'].upper(),
                    type_title,
                    context['year'],
                    context['user'].username,
                    context['today'].strftime("%d/%m/%Y - %H:%M:%S")
                ),
                merge_format
            )
            worksheet.merge_range(
                "A4:M4",
                context['title_sale_types'],
                merge_format_left
            )
            worksheet.merge_range(
                "A5:M5",
                context['title_hotels'],
                merge_format_left
            )
        else:
            columns = (len(context['total_by_dates'])-1) if len(context['total_by_dates']) > 12 else 11
            int_div = (columns)//2
            worksheet.merge_range(
                0,0,2,int_div, 
                "ON TOUR BY RCD - {}".format(self.property.name[4:]),
                merge_format
            )
            worksheet.merge_range(
                0,int_div+1, 2, columns+1,
                "RESUMEN DE {} DEL {} AL {} \n Reporte impreso por: {} el {}".format(
                    context['title'].upper(),
                    context['start_date'].strftime("%d/%m/%Y"),
                    context['due_date'].strftime("%d/%m/%Y"),
                    context['user'].username,
                    context['today'].strftime("%d/%m/%Y - %H:%M:%S")
                ),
                merge_format
            )
            worksheet.merge_range(
                3,0, 3,columns+1,
                context['title_sale_types'],
                merge_format_left
            )
            worksheet.merge_range(
                4,0, 4,columns+1,
                context['title_hotels'],
                merge_format_left
            )
        worksheet.write("A6", "TIPO DE VENTA" if context['sort_by'] == "sales_type" else "HOTEL", header_format)
        worksheet.set_column('A:A', 25)
        column = 1
        for date in context['total_by_dates']:
            if self.year is not None:
                worksheet.write(5,column, months[date['date']-1], header_format)
            else:
                worksheet.write(5,column, date['date'].strftime("%d/%m"), header_format)
            worksheet.set_column(column,column, 10)
            column += 1
        row = 6
        for reservation_services_group in context['reservation_services_groups']:
            worksheet.write(row,0, reservation_services_group['name'])
            column = 1
            for date in reservation_services_group['dates']:
                worksheet.write(row,column, "{}.{}".format(
                    str(int(date['adults_total'] if date['adults_total'] is not None else 0)),
                    str(int(date['childs_total'] if date['childs_total'] is not None else 0))))
                column += 1
            row += 1
        column = 1
        worksheet.write(row,0, "", unit_format)
        for total_date in context['total_by_dates']:
            worksheet.write(row,column, "{}.{}".format(
                str(int(total_date['adults_total'] if total_date['adults_total'] is not None else 0)),
                str(int(total_date['childs_total'] if total_date['childs_total'] is not None else 0))
            ),unit_format)
            column += 1
        if context['print_total'] is True:
            worksheet = workbook.add_worksheet('Totales')
            worksheet.set_column('A:A', 20)
            worksheet.set_column('B:B', 20)
            row = 0
        else:
            row += 1
        if self.year is not None:
            worksheet.merge_range(
                row,0,row,1, 
                "TOTAL DE PAX {} EN EL AÑO {}".format(
                    context['title'].upper(),
                    context['year'],
                ),
                total_format
            )
            worksheet.merge_range(
                row,4,row+1,10, 
                "EL TOTAL DE PAX QUE {} EN EL AÑO {} DE {}.{}".format(
                    "SALIERON" if context['type'] == "sales_type" else "LLEGARON", 
                    context['year'],
                    str(int(context['adults_total'] if context['adults_total'] is not None else 0)),
                    str(int(context['childs_total'] if context['childs_total'] is not None else 0))
                ),
                total_format_
            )
        else:
            worksheet.merge_range(
                row,0,row,1, 
                "TOTAL DE PAX {} DEL {} AL {}".format(
                    context['title'].upper(),
                    context['start_date'].strftime("%d/%m/%Y"),
                    context['due_date'].strftime("%d/%m/%Y"),
                ),
                total_format
            )
            worksheet.merge_range(
                row,4,row+1,10, 
                "EL TOTAL DE PAX QUE {} EN EL PERIODO COMPRENDIDO DEL {} AL DE {}.{}".format(
                    "SALIERON" if context['type'] == "sales_type" else "LLEGARON", 
                    context['start_date'].strftime("%d/%m/%Y"),
                    context['due_date'].strftime("%d/%m/%Y"),
                    str(int(context['adults_total'] if context['adults_total'] is not None else 0)),
                    str(int(context['childs_total'] if context['childs_total'] is not None else 0))
                ),
                total_format_
            )
        row_ = row + 1
        for reservation_services_group in context['reservation_services_groups']:
            worksheet.write(row_,0, reservation_services_group['name'])
            worksheet.write(row_,1, "{}.{}".format(
                str(int(reservation_services_group['adults_total'] if reservation_services_group['adults_total'] is not None else 0)),
                str(int(reservation_services_group['childs_total'] if reservation_services_group['childs_total'] is not None else 0)))
            )
            row_ += 1
        workbook.close()
        output.seek(0)
        if self.year is not None:
            filename = "{}-operation-report-summary-{}-{}.xlsx".format(str(self.year),self.property.code,str(datetime.now()))
        else:
            filename = "{}-{}-operation-report-summary-{}-{}.xlsx".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        response = HttpResponse(
            output.getvalue(),
            content_type='application/ms-excel;charset=utf-8'
        )
        response['Content-Disposition'] = "attachment; filename={}".format(filename)
        response["Content-Encoding"] = "UTF-8"
        return response
    
    def report_in_pdf(self,context):
        if self.year is not None:
            template = get_template('operation_report_summary_by_year.html')
            filename = "{}-operation-report-summary-{}-{}.pdf".format(str(self.year),self.property.code,str(datetime.now()))
        else:
            template = get_template('operation_report_summary_by_dates.html')
            filename = "{}-{}-operation-report-summary-{}-{}.pdf".format(str(self.start_date),str(self.due_date),self.property.code,str(datetime.now()))
        html  = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            
            response = HttpResponse(
                result.getvalue(), 
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = "inline; filename={}".format(filename)
            return response
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)

class InvoiceOperationReportController():
    def __init__(self, request, start_date, due_date, provider, format):
        self.request = request
        self.start_date = start_date
        self.due_date = due_date
        self.provider = provider
        self.format = format

    def filters(self):
        reservations_services = models.ReservationService.objects.dateRange(self.start_date,self.due_date).bySameProvider(self.provider).distinct().tableOptimization()
        return reservations_services
    
    def get_context(self):
        reservations_services = self.filters()
        if self.format == "pdf":
            return self.report_in_pdf(reservations_services)
        elif self.format == "excel":
            return self.report_in_excel(reservations_services)

    def report_in_excel(self, reservations_services):
        from SalesApp.controllers  import sale_subtotal
        import locale
        locale.setlocale( locale.LC_ALL, 'en_US.UTF-8' )
        output = BytesIO()
        filename = 'operation-invoice-report-{}.xlsx'.format(str(datetime.now()))
        workbook = xlsxwriter.Workbook(output)
        merge_format = workbook.add_format({
            "bold": 1,
            "align": "center",
            "valign": "vcenter",
        })
        header_format = workbook.add_format({
            'bold': True,
            'top':1,
            'bottom':1,
            "align": "center",
            "valign": "vcenter",
            'font_size':15,
        })
        id_format = workbook.add_format({
            "align": "center",
            "valign": "vcenter",
        })
        unit_format = workbook.add_format({
            "align": "right",
            "valign": "vcenter",
        })
        worksheet = workbook.add_worksheet()
        worksheet.merge_range("A1:L1", "REPORTE DE FACTURAS DE PROVEEDOR - {}".format(self.provider.name), merge_format)
        worksheet.merge_range("A2:L2", "DEL {} AL {}".format(self.start_date.strftime("%d/%m/%Y"),self.due_date.strftime("%d/%m/%Y")), merge_format)
        worksheet.write("A3", "Fecha", header_format)
        worksheet.set_column('A:A', 25)
        worksheet.write("B3", "referencia#", header_format)
        worksheet.set_column('B:D', 20)
        worksheet.write("C3", "#Cupon", header_format)
        worksheet.write("D3", "#Boarding Pass", header_format)
        worksheet.write("E3", "Servicio", header_format)
        worksheet.set_column('E:E', 20)
        worksheet.write("F3", "Pax", header_format)
        worksheet.set_column('F:F', 5)
        worksheet.write("G3", "Costo USD", header_format)
        worksheet.set_column('G:J', 15)
        worksheet.write("H3", "Costo MN", header_format)
        worksheet.write("I3", "Efectivo", header_format)
        worksheet.write("J3", "No Show", header_format)
        row = 3
        for reservation_service in reservations_services:
            worksheet.write(row,0, reservation_service.date.strftime("%d/%m/%Y"))
            worksheet.write(row,1, str(reservation_service.reservation.id).zfill(6),id_format)
            worksheet.write(row,3, str(reservation_service.id).zfill(8),id_format)
            transfers = dict(TRANSFER_TYPE_CHOICES_SIMPLE)
            worksheet.write(row,4,"{}-{}".format(reservation_service.service.name,transfers[reservation_service.transfer_type]))
            worksheet.write(row,5, reservation_service.adults+reservation_service.childs)
            cost_usd = 0
            cost_currency = 0
            sale = reservation_service.sales.filter(status="A").first()
            if sale is not None:
                worksheet.write(row,2, str(sale.sale_key).zfill(8),id_format)
                sale_total = sale_subtotal(sale)
                cost_usd = sale_total['total_cost_num']
                worksheet.write(row,6, "{} USD".format(locale.currency(cost_usd,grouping=True)),unit_format)
            else:
                worksheet.write(row,2, "")
                worksheet.write(row,6, "0 USD",unit_format)
            provider = reservation_service.service.provider
            if provider.currency == "MN":
                exchange_rate = ExchangeRate.objects.filter(start_date__lte=reservation_service.date,provider=provider).order_by("-start_date").first()
                if exchange_rate is not None:
                    cost_currency = cost_usd*exchange_rate.usd_currency
            worksheet.write(row,7, "{} {}".format(locale.currency(cost_currency,grouping=True),provider.currency),unit_format)
            worksheet.write(row,8, "X" if reservation_service.invoice_type == "cash" else "",id_format)
            worksheet.write(row,9, "X" if reservation_service.invoice_type == "no_show" else "",id_format)
            row += 1
        workbook.close()
        output.seek(0)
        response = HttpResponse(
            output.getvalue(),
            content_type='application/ms-excel;charset=utf-8'
        )
        response['Content-Disposition'] = "attachment; filename={}".format(filename)
        response["Content-Encoding"] = "UTF-8"
        return response
    
    def report_in_pdf(self,reservations_services):
        from SalesApp.controllers  import sale_subtotal
        template = get_template('operation_report_invoice.html')
        filename = 'operation-invoice-report-{}.pdf'.format(str(datetime.now()))
        context = {
            'host'                              :   self.request.get_host(),
            'reservations_services'             :   [],
            'provider'                          :   self.provider,
            'start_date'                        :   self.start_date,
            'due_date'                          :   self.due_date,
        }
        for reservations_service in reservations_services:
            reservations_service_data = serializers.ReservationServiceInvoiceSerializer(reservations_service).data
            transfers = dict(TRANSFER_TYPE_CHOICES_SIMPLE)
            reservations_service_data['date'] = reservations_service.date.strftime("%d/%m/%Y")
            reservations_service_data['transfer_type_data'] = "{}-{}".format(reservations_service.service.name,transfers[reservations_service.transfer_type])
            reservations_service_data['pax_data'] = reservations_service.adults+reservations_service.childs
            sale = reservations_service.sales.filter(status="A").first()
            if sale is not None:
                reservations_service_data['coupon'] = str(sale.sale_key).zfill(8)
                sale_total = sale_subtotal(sale)
                reservations_service_data['cost_usd'] = sale_total['total_cost_num']
                reservations_service_data['cost_currency'] = 0
            else:
                reservations_service_data['coupon'] = None
                reservations_service_data['cost_usd'] = 0
                reservations_service_data['cost_currency'] = 0
            provider = reservations_service.service.provider
            if provider.currency == "MN":
                exchange_rate = ExchangeRate.objects.filter(start_date__lte=reservations_service.date,provider=provider).order_by("-start_date").first()
                if exchange_rate is not None:
                    reservations_service_data['cost_currency'] = reservations_service_data['cost_usd']*exchange_rate.usd_currency
            context['reservations_services'].append(reservations_service_data)
        html  = template.render(context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("latin-1",'replace')), result)
        if not pdf.err:
            response = HttpResponse(
                result.getvalue(), 
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = "inline; filename={}".format(filename)
            return response
        raise CustomValidation(pdf.err, 'document', status.HTTP_400_BAD_REQUEST)

